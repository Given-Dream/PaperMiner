#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PaperMiner - 智能论文内容提取工具
功能：
1. 批量处理 input 文件夹中的所有 PDF
2. 自动提取文字、公式、图片、表格
3. 智能识别图片编号（Fig 1, Figure 2 等）
4. 智能章节提取（正则表达式 + LLM质量检查）
5. 生成规范的输出目录结构（extract文件夹）
"""

import tkinter as tk
from tkinter import filedialog, messagebox

try:
    import ttkbootstrap as ttk
except ImportError as exc:
    ttk = None
    _TTKBOOTSTRAP_IMPORT_ERROR = exc
else:
    _TTKBOOTSTRAP_IMPORT_ERROR = None
import subprocess
import threading
import queue
import sys
import io
import time
import os
import atexit
import faulthandler
import gc
import traceback
from pathlib import Path
import json
import re
import shutil
from typing import List

_STARTUP_MESSAGES = []

# PaperMiner.exe 使用 pythonw.exe 直接启动本脚本。运行环境在导入 MinerU
# 或 PyTorch 前完成初始化，不再依赖 PowerShell 修改环境变量。
_PAPERMINER_DIRECT_LAUNCH = "--paperminer-launcher" in sys.argv
_RUNTIME_PATH = Path(sys.executable).resolve().parent
_RUNTIME_PATH_ENTRIES = [
    _RUNTIME_PATH,
    _RUNTIME_PATH / "Scripts",
    _RUNTIME_PATH / "Library" / "bin",
    _RUNTIME_PATH / "Library" / "usr" / "bin",
]
os.environ["PATH"] = os.pathsep.join(
    [str(path) for path in _RUNTIME_PATH_ENTRIES] + [os.environ.get("PATH", "")]
)
os.environ.setdefault("KMP_DUPLICATE_LIB_OK", "TRUE")
os.environ.setdefault("OMP_NUM_THREADS", "1")
os.environ.setdefault("MINERU_MODEL_SOURCE", "modelscope")

# 导入版本信息
try:
    from version import __version__, __app_name__, __contact_email__
except ImportError:
    __version__ = "1.4.5"
    __app_name__ = "PaperMiner"
    __contact_email__ = "2878705044@qq.com"

# 导入 LLM 辅助模块
try:
    from llm_helper import LLMHelper, load_prompt_template, save_sections
    from llm_config import (
        DEEPSEEK_API_BASE,
        DEFAULT_DEEPSEEK_MODEL,
        LLMSettings,
        discover_models,
        format_speed_result,
        load_llm_settings,
        normalize_api_base_url,
        save_llm_settings,
        test_model_speed,
    )
    LLM_AVAILABLE = True
except ImportError as exc:
    LLM_AVAILABLE = False
    _STARTUP_MESSAGES.append(
        f"警告: 无法导入 LLM 模块，章节提取功能将不可用: {exc}"
    )

try:
    from section_merger import (
        merge_all_sections_and_charts_to_markdown as merge_sections_and_charts_to_markdown_files,
    )
except ImportError:
    from scripts.section_merger import (
        merge_all_sections_and_charts_to_markdown as merge_sections_and_charts_to_markdown_files,
    )

# 设置标准输出编码为 UTF-8。
# pythonw.exe 没有控制台，sys.stdout / sys.stderr 会是 None；doclayout_yolo
# 在导入时会直接读取 sys.stdout.encoding，因此必须在导入 MinerU 之前补建
# 一个完整的文本流。os.devnull 返回的 TextIOWrapper 同时支持 encoding、
# write、flush、fileno 和 reconfigure，能兼容 logging 及第三方模型库。
def _prepare_standard_stream(stream_name: str):
    stream = getattr(sys, stream_name, None)
    if stream is None or getattr(stream, "encoding", None) is None:
        stream = open(os.devnull, "w", encoding="utf-8", errors="replace")
        setattr(sys, stream_name, stream)
        return stream

    if sys.platform == "win32" and stream.encoding.lower().replace("-", "") != "utf8":
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except (AttributeError, OSError, ValueError):
            if hasattr(stream, "buffer"):
                stream = io.TextIOWrapper(
                    stream.buffer,
                    encoding="utf-8",
                    errors="replace",
                    line_buffering=True,
                )
                setattr(sys, stream_name, stream)
    return stream


_prepare_standard_stream("stdout")
_prepare_standard_stream("stderr")


# ---- LaTeX 数学片段 -> Unicode 纯文本（用于表格单元 / caption）----
# mineru 3.1.0 在 table_body HTML 里把上下标、希腊字母、科学计数法都保留为 LaTeX 源码
# （例如 "$\sigma ( \mathsf { S } \mathsf { c m } ^ { - 1 } )$"），下游的 openpyxl
# 写出去会是一片 "\sigma ... \times ..." 看起来像乱码。这个函数负责把 $...$ 段转成
# 可读的 Unicode：希腊字母、上下标、\mathrm/\mathsf 等 wrapper 去掉。
_SUP_TABLE = str.maketrans(
    "0123456789+-=().abcdefghijklmnoprstuvwxyz",
    "⁰¹²³⁴⁵⁶⁷⁸⁹"
    "⁺⁻⁼⁽⁾·"
    "ᵃᵇᶜᵈᵉᶠᵍʰⁱʲ"
    "ᵏˡᵐⁿᵒᵖʳˢᵗᵘᵛʷˣʸᶻ"
)
_SUB_TABLE = str.maketrans(
    "0123456789+-=().aehijklmnoprstuvx",
    "₀₁₂₃₄₅₆₇₈₉"
    "₊₋₌₍₎·"
    "ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ"
)

# ^{\circ}、^{\prime} 这种"伪上下标"实际是独立符号（度数、prime 等），不是字母上标。
# 如果按字母表 translate 会得到畸形字符串，之后又被 pylatexenc 当未知命令吃掉。
# 这里直接在 _to_sup_unicode / _to_sub_unicode 里识别，返回对应 Unicode 字面量。
_LATEX_SYMBOL_PASSTHROUGH = {
    r"\circ": "°",
    r"\prime": "′",
    r"\ast": "∗",
    r"\star": "⋆",
    r"\cdot": "·",
}

# 单词保护占位：用 NBSP 把 `{\bf of}` 这类 switch 形式内部的字母间空格保住，
# 避开下面 99 行的"两字母间去空格"，结尾再统一还原成普通空格。
_NBSP = " "


def _to_sup_unicode(s: str) -> str:
    inner = re.sub(r"\s+", "", s)
    if inner in _LATEX_SYMBOL_PASSTHROUGH:
        return _LATEX_SYMBOL_PASSTHROUGH[inner]
    try:
        return inner.translate(_SUP_TABLE)
    except Exception:
        return "^(" + inner + ")"


def _to_sub_unicode(s: str) -> str:
    inner = re.sub(r"\s+", "", s)
    if inner in _LATEX_SYMBOL_PASSTHROUGH:
        return _LATEX_SYMBOL_PASSTHROUGH[inner]
    try:
        return inner.translate(_SUB_TABLE)
    except Exception:
        return "_(" + inner + ")"


def _latex_math_segment_to_text(src: str) -> str:
    s = src.strip()
    for cmd in ("mathrm", "mathsf", "mathbf", "mathit", "mathfrak", "mathcal", "mathbb", "text"):
        s = re.sub(r"\\" + cmd + r"\s*\{\s*([^{}]*)\s*\}", r"\1", s)
        # mineru 偶发输出不带花括号的单字符形式 "\mathsf C"；若不展平，pylatexenc
        # 会渲染成数学 Astral 字母（如 U+1D5A2），Word 字体支持差会显示成方块。
        s = re.sub(r"\\" + cmd + r"\s+([A-Za-z])(?![A-Za-z])", r"\1", s)

    # Switch 形式 {\bf ...}、{\it ...} 常用于文本模式加粗/斜体；mineru 抽取时常
    # 把单词按字符拆开成 "{\bf o f}"。这里判断：纯字母序列视为单词，用 NBSP 包住
    # 两侧以扛过后续空格压缩；含数字/符号的（如 "{\bf P 1}"、"{\bf - P 5}"）按
    # 原内容展开，让常规的字母间去空格规则自然处理。
    def _switch_repl(m: "re.Match[str]") -> str:
        inner = m.group(1).strip()
        if re.fullmatch(r"[A-Za-z]+(?:\s+[A-Za-z]+)*", inner):
            word = re.sub(r"\s+", "", inner)
            return _NBSP + word + _NBSP
        return inner

    for sw in ("bf", "it", "rm", "sf", "tt"):
        s = re.sub(r"\{\s*\\" + sw + r"\s+([^{}]*?)\s*\}", _switch_repl, s)

    s = re.sub(r"\^\s*\{([^{}]*)\}", lambda m: _to_sup_unicode(m.group(1)), s)
    s = re.sub(r"_\s*\{([^{}]*)\}", lambda m: _to_sub_unicode(m.group(1)), s)
    s = re.sub(r"\^\s*([0-9A-Za-z+\-])", lambda m: _to_sup_unicode(m.group(1)), s)
    s = re.sub(r"_\s*([0-9A-Za-z+\-])", lambda m: _to_sub_unicode(m.group(1)), s)
    try:
        from pylatexenc.latex2text import LatexNodes2Text
        s = LatexNodes2Text().latex_to_text(s)
    except Exception:
        pass
    s = s.replace("{", "").replace("}", "")
    # 仅压缩 ASCII 空白，保留 NBSP 以维持单词边界保护。
    s = re.sub(r"[ \t\r\n\f\v]+", " ", s).strip()
    s = re.sub(r"(?<=[0-9A-Za-z.]) (?=[0-9A-Za-z.])", "", s)
    s = re.sub(r"\s+([)\]}])", r"\1", s)
    s = re.sub(r"([(\[{])\s+", r"\1", s)
    # "° C" → "°C"（数字侧的 "30 °C" 保留不动，符合科技写作惯例）
    s = re.sub(r"°\s+(?=[A-Za-z])", "°", s)
    # 还原单词保护符，再压一次可能出现的连续空格。
    s = s.replace(_NBSP, " ")
    s = re.sub(r" {2,}", " ", s).strip()
    return s


def latex_to_unicode(s):
    if not isinstance(s, str) or "$" not in s:
        return s
    return re.sub(r"\$([^$]*)\$", lambda m: _latex_math_segment_to_text(m.group(1)), s)


# ---- PDF 区域重渲染（用于从 mineru 3.x 的子面板 bbox 重建完整 Figure 图）----
# mineru 3.1.0 把多面板 Figure 切成 N 个子面板块，每块有独立 bbox。旧版 2.5.3 则
# 每个 Figure 就是一张完整图。我们的做法：同一页上所有 image/chart 块的 bbox 求并集，
# 从原 PDF 重渲染那个区域还原完整 Figure。
def _union_bbox(bboxes):
    """Union a list of (x1,y1,x2,y2) tuples. Returns None if empty."""
    if not bboxes:
        return None
    xs1 = [b[0] for b in bboxes]
    ys1 = [b[1] for b in bboxes]
    xs2 = [b[2] for b in bboxes]
    ys2 = [b[3] for b in bboxes]
    return (min(xs1), min(ys1), max(xs2), max(ys2))


def _repair_docx_missing_parts(docx_path):
    """自愈 python-docx 偶发 bug：保存后 [Content_Types].xml / rels 声明了某 XML 部件，
    但 zip 里该文件缺失，导致 Word 打开报"内容有问题"。

    做两件事：
    1. 扫描 [Content_Types].xml 里 `<Override PartName=...>` 列出的每个 xml 部件，
       若 zip 里对应文件缺失，注入一个最小合法的同名 XML 占位（对常见 WordML 部件
       内置模板；其它情况写一个最小根元素）。
    2. 扫描 word/_rels/document.xml.rels 里 Target 指向的相对路径，若 zip 里没有
       目标文件，同样补一个最小占位。

    这样不会覆盖已有内容，只在缺失时兜底。
    """
    import zipfile
    import shutil
    from pathlib import Path as _Path

    docx_path = _Path(docx_path)
    if not docx_path.exists():
        return

    try:
        with zipfile.ZipFile(docx_path, 'r') as zr:
            members = set(zr.namelist())
            try:
                ct_xml = zr.read('[Content_Types].xml').decode('utf-8', errors='replace')
            except KeyError:
                return
            try:
                rels_xml = zr.read('word/_rels/document.xml.rels').decode('utf-8', errors='replace')
            except KeyError:
                rels_xml = ''
    except Exception:
        return

    import re as _re

    # 最小合法的 WordML 部件模板（命名空间与 Word 要求一致）。未列出的用通用 root。
    minimal_parts = {
        'numbering': (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        ),
        'footnotes': (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        ),
        'endnotes': (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        ),
        'comments': (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        ),
    }

    def _pick_template(part_name):
        for key, body in minimal_parts.items():
            if key in part_name.lower():
                return body
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<root/>'
        )

    missing = []
    # 从 Content_Types 抓声明的 xml Override 部件
    for m in _re.finditer(
        r'<Override\s+[^>]*PartName="([^"]+\.xml)"', ct_xml
    ):
        part = m.group(1).lstrip('/')
        if part not in members:
            missing.append(part)

    # 从 document.xml.rels 抓声明的相对 target（转成 zip 内绝对路径）
    if rels_xml:
        for m in _re.finditer(
            r'<Relationship\s+[^/]*Target="([^"]+)"', rels_xml
        ):
            target = m.group(1)
            # Rels 在 word/_rels/document.xml.rels，target 相对 word/
            if target.startswith('/'):
                part = target.lstrip('/')
            elif target.startswith('..'):
                part = _re.sub(r'^\.\./', '', target)
            else:
                part = 'word/' + target
            # 跳过图片/字体等二进制（有 Default Extension 兜底）
            if not part.lower().endswith('.xml'):
                continue
            if part not in members and part not in missing:
                missing.append(part)

    if not missing:
        return

    # 用临时文件 + 追加写回，避免破坏原 zip 结构
    tmp = docx_path.with_suffix(docx_path.suffix + '.tmp')
    try:
        shutil.copy2(docx_path, tmp)
        with zipfile.ZipFile(tmp, 'a', zipfile.ZIP_DEFLATED) as zw:
            for part in missing:
                body = _pick_template(part)
                zw.writestr(part, body)
        shutil.move(str(tmp), str(docx_path))
    except Exception:
        if tmp.exists():
            tmp.unlink(missing_ok=True)


def _render_pdf_region(pdf_path, page_idx, bbox_pt, dpi=200, pad_pt=4):
    """Render a PDF page region to a PIL.Image.

    Args:
        pdf_path: path to the PDF file.
        page_idx: zero-based page index.
        bbox_pt: (x1,y1,x2,y2) in PDF points, top-left origin (same as mineru bbox).
        dpi: output resolution.
        pad_pt: padding in PDF points added around the bbox before cropping.
    """
    import pypdfium2 as pdfium
    doc = pdfium.PdfDocument(str(pdf_path))
    try:
        page = doc[page_idx]
        scale = dpi / 72.0
        pil_img = page.render(scale=scale).to_pil()
        x1, y1, x2, y2 = bbox_pt
        w_px, h_px = pil_img.size
        px_box = (
            max(0, int((x1 - pad_pt) * scale)),
            max(0, int((y1 - pad_pt) * scale)),
            min(w_px, int((x2 + pad_pt) * scale)),
            min(h_px, int((y2 + pad_pt) * scale)),
        )
        return pil_img.crop(px_box)
    finally:
        doc.close()


class BatchPDFProcessorGUI:
    def __init__(self, root):
        self.root = root
        self.root.title(f"{__app_name__} v{__version__} - 智能论文内容提取工具")
        # 横向看板优先利用宽屏，并在较小屏幕上自动收缩到可用区域。
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        self.initial_width = min(1500, max(900, screen_width - 80))
        self.initial_height = min(900, max(680, screen_height - 140))
        window_x = max(0, (screen_width - self.initial_width) // 2)
        window_y = max(0, (screen_height - self.initial_height) // 2)
        self.root.geometry(
            f"{self.initial_width}x{self.initial_height}+{window_x}+{window_y}"
        )

        # 安装目录只保存程序文件；输入/输出目录由用户配置决定并持久化到
        # %LOCALAPPDATA%\PaperMiner\settings.json，升级或改装程序时不会丢失。
        self.base_path = Path(__file__).parent.parent
        local_app_data = Path(
            os.environ.get("LOCALAPPDATA", str(Path.home() / "AppData" / "Local"))
        )
        self.user_config_path = local_app_data / "PaperMiner" / "settings.json"
        self.input_path = self.base_path / "input"
        self.output_path = self.base_path / "output"
        self.raw_output_path = self.output_path / "raw"
        self.extract_output_path = self.output_path / "extract"
        self._load_directory_preferences()

        # 正常运行由 PaperMiner.exe 直接启动 pythonw.exe，不再依赖外部
        # PowerShell 窗口。所有应用日志同时进入 GUI 和 UTF-8 日志文件。
        self._log_lock = threading.Lock()
        self._log_handle = None
        self.log_file_path = None
        self._fault_handler_enabled = False
        self._initialize_file_logging()

        # Tk 只能由主线程操作。后台处理和 loguru 回调只向此队列投递任务，
        # 主线程定时消费，避免长批次中的跨线程 Tcl 调用导致随机退出。
        self._ui_queue = queue.Queue()
        self._closing = False
        self._active_run_options = {}

        # 每次启动都从 .env 读取接口、已启用模型和当前模型。
        self.llm_settings = (
            load_llm_settings(self.base_path / ".env") if LLM_AVAILABLE else None
        )

        # 处理状态
        self.is_processing = False
        self.current_pdf_index = 0
        self.total_pdfs = 0
        self.success_count = 0
        self.failed_count = 0
        self.skipped_count = 0
        self.last_mineru_issue_code = None
        self._mineru_process = None
        # 颜色由 ttkbootstrap 主题接管；这些值同时供原生 Text/Toplevel 使用。
        self.bg_color = '#F4F7FB'
        self.card_bg = '#FFFFFF'
        self.fg_color = '#172033'
        self.accent_color = '#2563EB'

        # 配置窗口：minsize 覆盖缩小场景下的布局底线，不再与 geometry 冲突
        self.root.minsize(min(1000, self.initial_width), min(680, self.initial_height))
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)

        # 配置样式
        self.setup_styles()

        self.create_widgets()
        self._install_exception_hooks()
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.after(50, self._drain_ui_queue)
        self.log(f"{__app_name__} v{__version__} 已启动")
        if _PAPERMINER_DIRECT_LAUNCH:
            self.log("启动方式: PaperMiner.exe -> pythonw.exe (无 PowerShell)")
            self.log(f"Python: {sys.executable}")
            self.log(f"MinerU 环境: {_RUNTIME_PATH}")
            self.log(f"模型源: {os.environ.get('MINERU_MODEL_SOURCE', 'modelscope')}")
        diagnostics = os.environ.get("PAPERMINER_LAUNCH_DIAGNOSTICS", "").strip()
        if diagnostics:
            for line in diagnostics.splitlines():
                if line.strip():
                    self.log(line.strip())
        if self.log_file_path is not None:
            self.log(f"日志文件: {self.log_file_path}")
        for startup_message in _STARTUP_MESSAGES:
            self.log(startup_message)
        self.check_input_folder()

        # 让 PanedWindow 完成首次布局；不再把窗口宽度重置为旧版的 960px。
        self.root.update_idletasks()

        # 首次运行环境检测（延迟执行，避免阻塞 UI 启动）
        self.root.after(500, self.check_environment)

    def _load_directory_preferences(self):
        """读取用户选择的输入/输出目录；损坏配置自动回退到安装目录默认值。"""
        try:
            if not self.user_config_path.exists():
                return
            payload = json.loads(self.user_config_path.read_text(encoding="utf-8"))
            if not isinstance(payload, dict):
                return

            input_value = str(payload.get("input_directory", "")).strip()
            output_value = str(payload.get("output_directory", "")).strip()
            if input_value:
                candidate = Path(os.path.expandvars(input_value)).expanduser()
                if candidate.is_absolute():
                    self.input_path = candidate
            if output_value:
                candidate = Path(os.path.expandvars(output_value)).expanduser()
                if candidate.is_absolute():
                    self.output_path = candidate
        except (OSError, ValueError, TypeError, json.JSONDecodeError) as exc:
            _STARTUP_MESSAGES.append(f"警告: 目录配置读取失败，已使用默认目录: {exc}")
        finally:
            self._update_output_paths()

    def _update_output_paths(self):
        """由输出根目录统一派生 raw 与 extract，避免不同功能各自硬编码。"""
        self.raw_output_path = self.output_path / "raw"
        self.extract_output_path = self.output_path / "extract"

    def _save_directory_preferences(self):
        """原子保存路径设置，避免异常退出留下半个 JSON 文件。"""
        try:
            self.user_config_path.parent.mkdir(parents=True, exist_ok=True)
            payload = {
                "schema_version": 1,
                "input_directory": str(self.input_path),
                "output_directory": str(self.output_path),
                "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
            }
            temp_path = self.user_config_path.with_suffix(".json.tmp")
            temp_path.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            temp_path.replace(self.user_config_path)
        except OSError as exc:
            self.log(f"[WARN] 无法保存目录配置: {exc}")

    def _refresh_directory_labels(self):
        if hasattr(self, "input_path_label"):
            self.input_path_label.config(text=str(self.input_path))
        if hasattr(self, "output_path_label"):
            self.output_path_label.config(text=str(self.output_path))

    def choose_input_directory(self):
        """选择 PDF 输入目录并立即刷新文件清单。"""
        if self.is_processing:
            messagebox.showwarning("任务进行中", "请先停止或等待当前任务结束，再更改输入目录。")
            return
        selected = filedialog.askdirectory(
            title="选择 PDF 输入目录",
            initialdir=str(self.input_path if self.input_path.exists() else self.base_path),
            mustexist=True,
            parent=self.root,
        )
        if not selected:
            return
        self.input_path = Path(selected)
        self._save_directory_preferences()
        self._refresh_directory_labels()
        self.log(f"输入目录已更改: {self.input_path}")
        self.check_input_folder()

    def choose_output_directory(self):
        """选择输出根目录；raw 与 extract 始终位于该目录下。"""
        if self.is_processing:
            messagebox.showwarning("任务进行中", "请先停止或等待当前任务结束，再更改输出目录。")
            return
        selected = filedialog.askdirectory(
            title="选择输出根目录",
            initialdir=str(self.output_path if self.output_path.exists() else self.base_path),
            mustexist=True,
            parent=self.root,
        )
        if not selected:
            return
        candidate = Path(selected)
        try:
            candidate.mkdir(parents=True, exist_ok=True)
            (candidate / "raw").mkdir(parents=True, exist_ok=True)
            (candidate / "extract").mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            messagebox.showerror("目录不可用", f"无法写入所选输出目录：\n{candidate}\n\n{exc}")
            return

        self.output_path = candidate
        self._update_output_paths()
        self._save_directory_preferences()
        self._refresh_directory_labels()
        self.log(f"输出目录已更改: {self.output_path}")

    def _post_ui(self, callback, *args, **kwargs):
        """线程安全地把 UI 操作投递给 Tk 主线程。"""
        if not self._closing:
            self._ui_queue.put((callback, args, kwargs))

    def _drain_ui_queue(self):
        """由 Tk 主线程消费日志和状态更新，单轮限量避免界面被日志淹没。"""
        if self._closing:
            return
        for _ in range(250):
            try:
                callback, args, kwargs = self._ui_queue.get_nowait()
            except queue.Empty:
                break
            try:
                callback(*args, **kwargs)
            except tk.TclError:
                if self._closing:
                    return
            except Exception as exc:
                # 单个状态控件更新失败不能终止整个 UI 消费循环。
                self.log(f"[ERROR] 界面状态更新失败: {exc}")
        try:
            self.root.after(50, self._drain_ui_queue)
        except tk.TclError:
            self._closing = True

    def _capture_run_options(self):
        """在主线程一次性读取 Tk 变量，后台线程只使用普通 Python 值。"""
        self._active_run_options = {
            "extract_text": bool(self.extract_text_var.get()),
            "extract_formula": bool(self.extract_formula_var.get()),
            "extract_figures": bool(self.extract_figures_var.get()),
            "extract_tables": bool(self.extract_tables_var.get()),
            "extract_sections": bool(self.extract_sections_var.get()),
            "use_gpu": bool(self.use_gpu_var.get()),
            "skip_processed": bool(self.skip_processed_var.get()),
            "backend": self.backend_var.get(),
            "llm_model": self.llm_model_var.get(),
            "llm_provider": self.llm_settings.provider if self.llm_settings else "deepseek",
        }

    def _run_option(self, name, default=None):
        return self._active_run_options.get(name, default)

    def _on_close(self):
        if self.is_processing and not messagebox.askyesno(
            "任务仍在运行",
            "当前 PDF 可能仍在 MinerU 中处理。强制关闭会中断本篇输出，确认退出吗？",
            parent=self.root,
        ):
            return
        self.is_processing = False
        self._terminate_active_mineru()
        self._closing = True
        self._close_log_file()
        self.root.destroy()

    def _terminate_active_mineru(self):
        """终止当前隔离进程；仅由用户点击停止或确认关闭时调用。"""
        process = self._mineru_process
        if process is None:
            return
        try:
            if process.poll() is None:
                process.terminate()
                self.log("已向当前 MinerU 隔离进程发送停止请求。")
        except (OSError, ProcessLookupError) as exc:
            self.log(f"停止 MinerU 隔离进程时出现提示: {exc}")

    def setup_styles(self):
        """配置 PaperMiner v1.4.x 的 ttkbootstrap 主题与少量品牌样式。"""
        style = getattr(self.root, 'style', None) or ttk.Style()
        self.style = style
        try:
            style.theme_use('bootstrap-light')
        except tk.TclError:
            pass

        colors = getattr(style, 'colors', None)
        if colors is not None:
            self.bg_color = getattr(colors, 'bg', self.bg_color)
            self.card_bg = getattr(colors, 'inputbg', self.card_bg)
            self.fg_color = getattr(colors, 'fg', self.fg_color)
            self.accent_color = getattr(colors, 'primary', self.accent_color)

        self.root.option_add('*Font', ('Microsoft YaHei UI', 9))
        self.root.configure(bg=self.bg_color)

        # 字体、间距与品牌区域。颜色型控件使用 bootstyle，避免重复硬编码。
        style.configure('TLabel', font=('Microsoft YaHei UI', 9))
        style.configure('TButton', font=('Microsoft YaHei UI', 9), padding=(12, 7))
        style.configure('TCheckbutton', font=('Microsoft YaHei UI', 9))
        style.configure('TRadiobutton', font=('Microsoft YaHei UI', 9))
        style.configure('TLabelframe.Label', font=('Microsoft YaHei UI', 10, 'bold'))
        style.configure('Hero.TFrame', background='#175CD3')
        style.configure(
            'HeroTitle.TLabel',
            background='#175CD3',
            foreground='#FFFFFF',
            font=('Microsoft YaHei UI', 16, 'bold'),
        )
        style.configure(
            'HeroSubtitle.TLabel',
            background='#175CD3',
            foreground='#DCE9FF',
            font=('Microsoft YaHei UI', 9),
        )
        style.configure(
            'HeroBadge.TLabel',
            background='#FFFFFF',
            foreground='#175CD3',
            font=('Segoe UI', 9, 'bold'),
            padding=(10, 5),
        )
        style.configure(
            'HeroAction.TButton',
            background='#FFFFFF',
            foreground='#175CD3',
            font=('Microsoft YaHei UI', 9, 'bold'),
            padding=(13, 7),
            borderwidth=0,
        )
        style.map(
            'HeroAction.TButton',
            background=[('active', '#EAF2FF'), ('pressed', '#DCE9FF')],
            foreground=[('active', '#1849A9'), ('pressed', '#1849A9')],
        )
        style.configure(
            'SectionTitle.TLabel',
            font=('Microsoft YaHei UI', 10, 'bold'),
            foreground='#172033',
        )
        style.configure(
            'Muted.TLabel',
            font=('Microsoft YaHei UI', 8),
            foreground='#667085',
        )
        style.configure(
            'Status.TLabel',
            font=('Microsoft YaHei UI', 13, 'bold'),
            foreground='#067647',
        )
        style.configure('SuccessStat.TFrame', background='#ECFDF3')
        style.configure('DangerStat.TFrame', background='#FEF3F2')
        style.configure('WarningStat.TFrame', background='#FFFAEB')
        for name, background, foreground in (
            ('Success', '#ECFDF3', '#067647'),
            ('Danger', '#FEF3F2', '#B42318'),
            ('Warning', '#FFFAEB', '#B54708'),
        ):
            style.configure(
                f'{name}StatTitle.TLabel',
                background=background,
                foreground='#475467',
                font=('Microsoft YaHei UI', 8),
            )
            style.configure(
                f'{name}StatValue.TLabel',
                background=background,
                foreground=foreground,
                font=('Segoe UI', 18, 'bold'),
            )

    def create_styled_checkbutton(self, parent, text, variable, bootstyle='primary'):
        """创建由 ttkbootstrap 统一渲染的复选框。"""
        return ttk.Checkbutton(
            parent,
            text=text,
            variable=variable,
            bootstyle=bootstyle,
        )

    def create_widgets(self):
        """创建 v1.4.x 横向工作台：配置、任务和实时日志始终同屏。"""
        main_frame = ttk.Frame(self.root, padding=14)
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(1, weight=1)

        # 品牌栏：版本与全局设置在同一视觉层级。
        hero = ttk.Frame(main_frame, style='Hero.TFrame', padding=(20, 12))
        hero.grid(row=0, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        hero.grid_columnconfigure(0, weight=1)
        hero_text = ttk.Frame(hero, style='Hero.TFrame')
        hero_text.grid(row=0, column=0, sticky=tk.W)
        ttk.Label(
            hero_text,
            text=__app_name__,
            style='HeroTitle.TLabel',
            font=('Microsoft YaHei UI', 16, 'bold'),
        ).pack(anchor=tk.W)
        ttk.Label(
            hero_text,
            text='智能论文内容提取工作台 · 配置、进度与日志同屏',
            style='HeroSubtitle.TLabel',
            font=('Microsoft YaHei UI', 9),
        ).pack(anchor=tk.W, pady=(2, 0))
        ttk.Label(hero, text=f'v{__version__}', style='HeroBadge.TLabel').grid(
            row=0, column=1, padx=(12, 10)
        )
        ttk.Button(
            hero,
            text='接口与模型设置',
            command=self.open_settings,
            style='HeroAction.TButton',
        ).grid(row=0, column=2)

        # ttkbootstrap 2.x 使用 Panedwindow；三栏都可由用户拖动分隔线调整。
        self.dashboard_paned = ttk.Panedwindow(main_frame, orient=tk.HORIZONTAL)
        self.dashboard_paned.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        config_board = ttk.Frame(self.dashboard_paned, padding=(0, 0, 7, 0))
        task_board = ttk.Frame(self.dashboard_paned, padding=(7, 0))
        log_board = ttk.Frame(self.dashboard_paned, padding=(7, 0, 0, 0))
        self.dashboard_paned.add(config_board, weight=36)
        self.dashboard_paned.add(task_board, weight=27)
        self.dashboard_paned.add(log_board, weight=37)

        config_board.grid_columnconfigure(0, weight=1)
        config_board.grid_rowconfigure(1, weight=1)
        task_board.grid_columnconfigure(0, weight=1)
        task_board.grid_rowconfigure(1, weight=1)
        log_board.grid_columnconfigure(0, weight=1)
        log_board.grid_rowconfigure(0, weight=1)

        # 左栏：输入文件与处理配置。
        info_frame = ttk.Labelframe(
            config_board, text='输入文件', padding=12, bootstyle='primary'
        )
        info_frame.grid(row=0, column=0, pady=(0, 9), sticky=(tk.W, tk.E))
        for column in range(3):
            info_frame.grid_columnconfigure(column, weight=1)
        self.file_count_label = ttk.Label(
            info_frame, text='PDF 文件数量: 0', style='SectionTitle.TLabel'
        )
        self.file_count_label.grid(row=0, column=0, columnspan=3, sticky=tk.W)
        self.input_path_label = ttk.Label(
            info_frame,
            text=str(self.input_path),
            style='Muted.TLabel',
            wraplength=410,
        )
        self.input_path_label.grid(
            row=1, column=0, columnspan=3, sticky=tk.W, pady=(3, 9)
        )
        ttk.Button(
            info_frame,
            text='刷新文件',
            command=self.check_input_folder,
            bootstyle='secondary outline',
        ).grid(row=2, column=0, sticky=(tk.W, tk.E), padx=(0, 3))
        ttk.Button(
            info_frame,
            text='选择目录',
            command=self.choose_input_directory,
            bootstyle='primary',
        ).grid(row=2, column=1, sticky=(tk.W, tk.E), padx=3)
        ttk.Button(
            info_frame,
            text='打开目录',
            command=lambda: self.open_folder(self.input_path),
            bootstyle='secondary outline',
        ).grid(row=2, column=2, sticky=(tk.W, tk.E), padx=(3, 0))

        options_frame = ttk.Labelframe(
            config_board, text='处理配置', padding=12, bootstyle='primary'
        )
        options_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        options_frame.grid_columnconfigure(0, weight=1)
        options_frame.grid_columnconfigure(1, weight=1)

        self.process_mode_var = tk.StringVar(value='full')
        ttk.Label(options_frame, text='处理模式', style='SectionTitle.TLabel').grid(
            row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 4)
        )
        ttk.Radiobutton(
            options_frame,
            text='完整流程',
            variable=self.process_mode_var,
            value='full',
            command=self.on_mode_change,
            bootstyle='primary',
        ).grid(row=1, column=0, sticky=tk.W, pady=2)
        ttk.Radiobutton(
            options_frame,
            text='仅提取 raw',
            variable=self.process_mode_var,
            value='extract_only',
            command=self.on_mode_change,
            bootstyle='primary',
        ).grid(row=1, column=1, sticky=tk.W, pady=2)
        ttk.Label(
            options_frame,
            text='完整流程：PDF → MinerU → 结构化提取',
            style='Muted.TLabel',
        ).grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(1, 7))
        ttk.Separator(options_frame, orient='horizontal', bootstyle='secondary').grid(
            row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 8)
        )

        self.extract_text_var = tk.BooleanVar(value=True)
        self.extract_formula_var = tk.BooleanVar(value=True)
        self.extract_figures_var = tk.BooleanVar(value=True)
        self.extract_tables_var = tk.BooleanVar(value=True)
        self.extract_sections_var = tk.BooleanVar(value=True)
        self.use_gpu_var = tk.BooleanVar(value=True)
        self.skip_processed_var = tk.BooleanVar(value=True)

        ttk.Label(options_frame, text='提取内容', style='SectionTitle.TLabel').grid(
            row=4, column=0, columnspan=2, sticky=tk.W, pady=(0, 4)
        )
        option_specs = (
            ('文字 / Markdown', self.extract_text_var, 5, 0),
            ('公式 / LaTeX', self.extract_formula_var, 5, 1),
            ('图片与编号', self.extract_figures_var, 6, 0),
            ('表格 / Excel', self.extract_tables_var, 6, 1),
        )
        for text, variable, row, column in option_specs:
            self.create_styled_checkbutton(options_frame, text, variable).grid(
                row=row, column=column, sticky=tk.W, pady=2
            )
        self.create_styled_checkbutton(
            options_frame,
            '论文章节（正则 + LLM）',
            self.extract_sections_var,
            bootstyle='info',
        ).grid(row=7, column=0, columnspan=2, sticky=tk.W, pady=2)

        ttk.Separator(options_frame, orient='horizontal', bootstyle='secondary').grid(
            row=8, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=8
        )

        llm_frame = ttk.Frame(options_frame)
        llm_frame.grid(row=9, column=0, columnspan=2, sticky=(tk.W, tk.E))
        llm_frame.grid_columnconfigure(1, weight=1)
        ttk.Label(llm_frame, text='LLM 模型').grid(
            row=0, column=0, sticky=tk.W, padx=(0, 8)
        )
        initial_models = self._configured_llm_models()
        initial_model = self.llm_settings.active_model if self.llm_settings else ''
        self.llm_model_var = tk.StringVar(value=initial_model)
        self.llm_model_combo = ttk.Combobox(
            llm_frame,
            textvariable=self.llm_model_var,
            values=initial_models,
            state='readonly' if initial_models else 'disabled',
            bootstyle='primary',
        )
        self.llm_model_combo.grid(row=0, column=1, sticky=(tk.W, tk.E))
        self.llm_model_combo.bind('<<ComboboxSelected>>', self._on_llm_model_selected)
        self.llm_provider_label = ttk.Label(
            llm_frame,
            text=f'接口：{self._llm_provider_display_name()}（在设置中管理）',
            style='Muted.TLabel',
            wraplength=410,
        )
        self.llm_provider_label.grid(
            row=1, column=0, columnspan=2, sticky=tk.W, pady=(4, 0)
        )

        backend_frame = ttk.Frame(options_frame)
        backend_frame.grid(
            row=10, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(9, 0)
        )
        backend_frame.grid_columnconfigure(1, weight=1)
        ttk.Label(backend_frame, text='MinerU 后端').grid(
            row=0, column=0, sticky=tk.W, padx=(0, 8)
        )
        self.backend_var = tk.StringVar(value='pipeline')
        self.backend_combo = ttk.Combobox(
            backend_frame,
            textvariable=self.backend_var,
            values=['pipeline', 'vlm-auto-engine', 'hybrid-auto-engine'],
            state='readonly',
            bootstyle='primary',
        )
        self.backend_combo.grid(row=0, column=1, sticky=(tk.W, tk.E))
        ttk.Label(
            backend_frame,
            text='pipeline 无需 VLM 模型，国内网络推荐',
            style='Muted.TLabel',
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, pady=(4, 0))

        toggle_frame = ttk.Frame(options_frame)
        toggle_frame.grid(
            row=11, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(10, 0)
        )
        toggle_frame.grid_columnconfigure(0, weight=1)
        toggle_frame.grid_columnconfigure(1, weight=1)
        self.gpu_checkbox = self.create_styled_checkbutton(
            toggle_frame,
            'GPU 加速',
            self.use_gpu_var,
            bootstyle='success round toggle',
        )
        self.gpu_checkbox.grid(row=0, column=0, sticky=tk.W)
        self.skip_checkbox = self.create_styled_checkbutton(
            toggle_frame,
            '跳过已有结果',
            self.skip_processed_var,
            bootstyle='info round toggle',
        )
        self.skip_checkbox.grid(row=0, column=1, sticky=tk.W)

        # 中栏：任务控制、进度统计和输出。
        control_frame = ttk.Labelframe(
            task_board, text='任务控制', padding=12, bootstyle='success'
        )
        control_frame.grid(row=0, column=0, pady=(0, 9), sticky=(tk.W, tk.E))
        control_frame.grid_columnconfigure(0, weight=2)
        control_frame.grid_columnconfigure(1, weight=1)
        self.start_button = ttk.Button(
            control_frame,
            text='开始处理',
            command=self.start_processing,
            bootstyle='success',
        )
        self.start_button.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 4))
        self.stop_button = ttk.Button(
            control_frame,
            text='停止',
            command=self.stop_processing,
            state='disabled',
            bootstyle='danger outline',
        )
        self.stop_button.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(4, 0))

        progress_frame = ttk.Labelframe(
            task_board, text='任务看板', padding=12, bootstyle='info'
        )
        progress_frame.grid(
            row=1, column=0, pady=(0, 9), sticky=(tk.W, tk.E, tk.N, tk.S)
        )
        progress_frame.grid_columnconfigure(0, weight=1)
        self.status_label = ttk.Label(
            progress_frame, text='准备就绪', style='Status.TLabel'
        )
        self.status_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 9))
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            mode='determinate',
            bootstyle='primary striped',
        )
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 8))
        self.progress_text = ttk.Label(
            progress_frame,
            text='等待任务',
            style='Muted.TLabel',
            wraplength=240,
        )
        self.progress_text.grid(row=2, column=0, sticky=tk.W, pady=(0, 13))

        stats_frame = ttk.Frame(progress_frame)
        stats_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))
        for col in range(3):
            stats_frame.grid_columnconfigure(col, weight=1)

        def stat_cell(col, label_text, style_prefix, pad):
            cell = ttk.Frame(stats_frame, style=f'{style_prefix}Stat.TFrame', padding=(8, 8))
            cell.grid(row=0, column=col, sticky=(tk.W, tk.E), padx=pad)
            ttk.Label(
                cell,
                text=label_text,
                style=f'{style_prefix}StatTitle.TLabel',
            ).pack()
            number = ttk.Label(
                cell,
                text='0',
                style=f'{style_prefix}StatValue.TLabel',
            )
            number.pack(pady=(1, 0))
            return number

        self.success_label = stat_cell(0, '成功', 'Success', (0, 3))
        self.failed_label = stat_cell(1, '失败', 'Danger', 3)
        self.skipped_label = stat_cell(2, '跳过', 'Warning', (3, 0))

        ttk.Label(
            progress_frame,
            text='执行详情会实时写入右侧日志及 logs 文件夹。',
            style='Muted.TLabel',
            wraplength=240,
        ).grid(row=4, column=0, sticky=tk.W, pady=(14, 0))

        output_frame = ttk.Labelframe(
            task_board, text='输出操作', padding=12, bootstyle='secondary'
        )
        output_frame.grid(row=2, column=0, sticky=(tk.W, tk.E))
        for column in range(3):
            output_frame.grid_columnconfigure(column, weight=1)
        self.output_path_label = ttk.Label(
            output_frame,
            text=str(self.output_path),
            style='Muted.TLabel',
            wraplength=240,
        )
        self.output_path_label.grid(
            row=0, column=0, columnspan=3, sticky=tk.W, pady=(0, 8)
        )
        ttk.Button(
            output_frame,
            text='选择输出',
            command=self.choose_output_directory,
            bootstyle='primary',
        ).grid(row=1, column=0, sticky=(tk.W, tk.E), padx=(0, 3))
        ttk.Button(
            output_frame,
            text='打开 raw',
            command=lambda: self.open_folder(self.raw_output_path),
            bootstyle='secondary outline',
        ).grid(row=1, column=1, sticky=(tk.W, tk.E), padx=3)
        ttk.Button(
            output_frame,
            text='打开 extract',
            command=lambda: self.open_folder(self.extract_output_path),
            bootstyle='secondary outline',
        ).grid(row=1, column=2, sticky=(tk.W, tk.E), padx=(3, 0))
        ttk.Button(
            output_frame,
            text='合并同名章节和图表到 Markdown',
            command=self.merge_all_sections_and_charts_to_markdown,
            bootstyle='info',
        ).grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(8, 0))

        # 右栏：常驻日志，不再依赖外部 PowerShell 窗口。
        log_frame = ttk.Labelframe(
            log_board, text='实时日志', padding=10, bootstyle='dark'
        )
        log_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.grid_columnconfigure(0, weight=1)
        log_frame.grid_rowconfigure(1, weight=1)
        log_toolbar = ttk.Frame(log_frame)
        log_toolbar.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 8))
        log_toolbar.grid_columnconfigure(0, weight=1)
        ttk.Label(
            log_toolbar,
            text='● 运行日志',
            bootstyle='success',
        ).grid(row=0, column=0, sticky=tk.W)

        def clear_log_view():
            self.log_text.config(state='normal')
            self.log_text.delete('1.0', tk.END)
            self.log_text.config(state='disabled')

        ttk.Button(
            log_toolbar,
            text='日志目录',
            command=lambda: self.open_folder(self.base_path / 'logs'),
            bootstyle='secondary outline',
        ).grid(row=0, column=1, padx=(6, 4))
        ttk.Button(
            log_toolbar,
            text='清空',
            command=clear_log_view,
            bootstyle='secondary outline',
        ).grid(row=0, column=2, padx=(4, 0))

        self.log_text = tk.Text(
            log_frame,
            wrap=tk.CHAR,
            font=('Cascadia Mono', 9),
            bg='#101828',
            fg='#E4E7EC',
            insertbackground='#FFFFFF',
            selectbackground='#344054',
            selectforeground='#FFFFFF',
            relief='flat',
            borderwidth=0,
            padx=12,
            pady=12,
            spacing1=2,
            spacing3=2,
            state='disabled',
        )
        self.log_text.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_scrollbar = ttk.Scrollbar(
            log_frame, orient='vertical', command=self.log_text.yview, bootstyle='dark round'
        )
        log_scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))
        self.log_text.config(yscrollcommand=log_scrollbar.set)

        footer_frame = ttk.Frame(main_frame)
        footer_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(8, 0))
        ttk.Label(
            footer_frame,
            text='拖动分隔线可调整配置 / 任务 / 日志宽度',
            style='Muted.TLabel',
        ).pack(side=tk.LEFT)
        ttk.Label(
            footer_frame,
            text=f'联系邮箱：{__contact_email__}',
            style='Muted.TLabel',
        ).pack(side=tk.RIGHT)

        self.root.after_idle(self._set_dashboard_sashes)

    def _set_dashboard_sashes(self):
        """首次显示时按 36% / 27% / 37% 分配三块看板。"""
        try:
            width = self.dashboard_paned.winfo_width()
            if width > 300:
                self.dashboard_paned.sashpos(0, int(width * 0.36))
                self.dashboard_paned.sashpos(1, int(width * 0.63))
        except tk.TclError:
            pass

    def _create_legacy_widgets(self):
        """创建界面组件"""
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(4, weight=1)

        # 标题区域 - 简约风格
        title_frame = tk.Frame(main_frame, bg=self.accent_color, height=44)
        title_frame.grid(row=0, column=0, pady=(0, 16), sticky=(tk.W, tk.E))
        title_frame.grid_propagate(False)

        title_label = tk.Label(
            title_frame,
            text=f"{__app_name__} v{__version__} - 智能论文内容提取工具",
            font=('Microsoft YaHei UI', 14, 'bold'),
            bg=self.accent_color,
            fg='white'
        )
        title_label.pack(side=tk.LEFT, expand=True, padx=(24, 0))

        # 设置按钮：白底蓝字，明显对比，避免被蓝色标题栏"吞掉"
        settings_button = tk.Button(
            title_frame,
            text="⚙ 设置",
            font=('Microsoft YaHei UI', 9),
            bg='#FFFFFF',
            fg=self.accent_color,
            activebackground='#E8EFFA',
            activeforeground=self.accent_color,
            relief='flat',
            borderwidth=0,
            cursor='hand2',
            padx=12,
            pady=2,
            command=self.open_settings
        )
        settings_button.pack(side=tk.RIGHT, padx=12, pady=8)
        
        # 文件信息区域
        info_frame = ttk.LabelFrame(main_frame, text="📄 文件信息", padding="12")
        info_frame.grid(row=1, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        info_frame.grid_columnconfigure(1, weight=1)

        self.file_count_label = ttk.Label(
            info_frame,
            text="PDF 文件数量: 0",
            font=('Microsoft YaHei UI', 10)
        )
        self.file_count_label.grid(row=0, column=0, sticky=tk.W, pady=4)

        # 显示当前监听的 input 路径（灰色小字），便于用户确认目录
        ttk.Label(
            info_frame,
            text=f"路径：{self.input_path}",
            font=('Microsoft YaHei UI', 8),
            foreground='#888'
        ).grid(row=0, column=1, sticky=tk.W, padx=(12, 8), pady=4)

        ttk.Button(
            info_frame,
            text="🔄 刷新",
            command=self.check_input_folder,
            width=10
        ).grid(row=0, column=2, sticky=tk.E, pady=4)

        ttk.Button(
            info_frame,
            text="📂 打开 input 文件夹",
            command=lambda: self.open_folder(self.input_path)
        ).grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=4)
        
        # 处理选项区域
        options_frame = ttk.LabelFrame(main_frame, text="⚙ 处理选项", padding="12")
        options_frame.grid(row=2, column=0, pady=(0, 12), sticky=(tk.W, tk.E))

        # 处理模式选择
        self.process_mode_var = tk.StringVar(value="full")

        mode_frame = ttk.Frame(options_frame)
        mode_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        ttk.Label(
            mode_frame,
            text="处理模式:",
            font=('Microsoft YaHei UI', 9)
        ).pack(side=tk.LEFT, padx=(0, 10))

        ttk.Radiobutton(
            mode_frame,
            text="完整处理 (PDF → MinerU → 提取)",
            variable=self.process_mode_var,
            value="full",
            command=self.on_mode_change
        ).pack(side=tk.LEFT, padx=(0, 15))

        ttk.Radiobutton(
            mode_frame,
            text="仅提取 (从已有 raw 文件夹提取)",
            variable=self.process_mode_var,
            value="extract_only",
            command=self.on_mode_change
        ).pack(side=tk.LEFT)

        # 分隔线
        ttk.Separator(options_frame, orient='horizontal').grid(
            row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=10
        )

        self.extract_text_var = tk.BooleanVar(value=True)
        self.extract_formula_var = tk.BooleanVar(value=True)
        self.extract_figures_var = tk.BooleanVar(value=True)
        self.extract_tables_var = tk.BooleanVar(value=True)
        self.extract_sections_var = tk.BooleanVar(value=True)  # 默认勾选
        self.use_gpu_var = tk.BooleanVar(value=True)
        self.skip_processed_var = tk.BooleanVar(value=True)

        # 使用统一样式的复选框（显示正确的勾选标记）
        self.create_styled_checkbutton(
            options_frame,
            "✏️ 提取文字 (Markdown)",
            self.extract_text_var
        ).grid(row=2, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "🔢 提取公式 (LaTeX)",
            self.extract_formula_var
        ).grid(row=3, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "🖼️提取图片(智能识别编号)",
            self.extract_figures_var
        ).grid(row=4, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "📊 提取表格 (Excel)",
            self.extract_tables_var
        ).grid(row=5, column=0, sticky=tk.W, pady=4)

        self.create_styled_checkbutton(
            options_frame,
            "📑 提取论文章节 (正则表达式 + LLM)",
            self.extract_sections_var
        ).grid(row=6, column=0, sticky=tk.W, pady=4)

        # 当前 LLM 接口及模型。自定义接口在设置中普通单击选择一个模型。
        llm_frame = ttk.Frame(options_frame)
        llm_frame.grid(row=7, column=0, sticky=tk.W, padx=(30, 0), pady=(5, 2))

        ttk.Label(llm_frame, text="🤖 LLM 模型:").pack(side=tk.LEFT, padx=(0, 5))

        initial_models = self._configured_llm_models()
        initial_model = self.llm_settings.active_model if self.llm_settings else ""
        self.llm_model_var = tk.StringVar(value=initial_model)
        self.llm_model_combo = ttk.Combobox(
            llm_frame,
            textvariable=self.llm_model_var,
            values=initial_models,
            state="readonly" if initial_models else "disabled",
            width=34,
        )
        self.llm_model_combo.pack(side=tk.LEFT, padx=(0, 8))
        self.llm_model_combo.bind("<<ComboboxSelected>>", self._on_llm_model_selected)

        provider_text = self._llm_provider_display_name()
        self.llm_provider_label = ttk.Label(
            llm_frame,
            text=f"接口：{provider_text}（在设置中管理）",
            font=('Microsoft YaHei UI', 8),
            foreground='#888888',
        )
        self.llm_provider_label.pack(side=tk.LEFT)

        ttk.Separator(options_frame, orient='horizontal').grid(
            row=8, column=0, sticky=(tk.W, tk.E), pady=10
        )

        # MinerU 后端选择
        backend_frame = ttk.Frame(options_frame)
        backend_frame.grid(row=9, column=0, sticky=tk.W, pady=4)

        ttk.Label(
            backend_frame,
            text="🔧 MinerU 后端:",
            font=('Microsoft YaHei UI', 9)
        ).pack(side=tk.LEFT, padx=(0, 5))

        self.backend_var = tk.StringVar(value="pipeline")
        self.backend_combo = ttk.Combobox(
            backend_frame,
            textvariable=self.backend_var,
            values=["pipeline", "vlm-auto-engine", "hybrid-auto-engine"],
            state="readonly",
            width=20
        )
        self.backend_combo.pack(side=tk.LEFT, padx=(0, 8))

        ttk.Label(
            backend_frame,
            text="(pipeline 无需 VLM 模型，推荐国内网络使用)",
            font=('Microsoft YaHei UI', 8),
            foreground='#888888'
        ).pack(side=tk.LEFT)

        self.gpu_checkbox = self.create_styled_checkbutton(
            options_frame,
            "⚡ 使用 GPU 加速 (推荐)",
            self.use_gpu_var
        )
        self.gpu_checkbox.grid(row=10, column=0, sticky=tk.W, pady=4)

        ttk.Separator(options_frame, orient='horizontal').grid(
            row=11, column=0, sticky=(tk.W, tk.E), pady=10
        )

        self.skip_checkbox = self.create_styled_checkbutton(
            options_frame,
            "跳过已处理的文件 (extract 中已有结果)",
            self.skip_processed_var
        )
        self.skip_checkbox.grid(row=12, column=0, sticky=tk.W, pady=4)

        # 控制按钮区域：开始：停止 = 2 : 1，体现主次（停止 90% 时间 disabled）
        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=3, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        control_frame.grid_columnconfigure(0, weight=2)
        control_frame.grid_columnconfigure(1, weight=1)

        self.start_button = ttk.Button(
            control_frame,
            text="开始处理",
            command=self.start_processing,
            style='Primary.TButton'
        )
        self.start_button.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 6))

        self.stop_button = ttk.Button(
            control_frame,
            text="停止",
            command=self.stop_processing,
            state='disabled',
            style='Stop.TButton'
        )
        self.stop_button.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=(6, 0))

        # 进度区域
        progress_frame = ttk.LabelFrame(main_frame, text="📊 处理进度", padding="12")
        progress_frame.grid(row=4, column=0, pady=(0, 12), sticky=(tk.W, tk.E, tk.N, tk.S))
        progress_frame.grid_columnconfigure(0, weight=1)
        # 日志行 weight=1 保证窗口变大时吸收剩余纵向空间；minsize=180 防止窗口
        # 变矮时日志区被压扁成一条"黑条"。
        progress_frame.grid_rowconfigure(4, weight=1, minsize=180)

        # 主行：状态 + "X / Y" 进度数字（处理中由 update_progress 覆写）
        self.status_label = ttk.Label(
            progress_frame,
            text="准备就绪",
            font=('Microsoft YaHei UI', 10, 'bold'),
            foreground='#27ae60'
        )
        self.status_label.grid(row=0, column=0, sticky=tk.W, pady=(0, 8))

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            mode='determinate'
        )
        self.progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 8))

        # 副行：显示当前处理的文件名（处理中时）。留空即不占视线。
        self.progress_text = ttk.Label(
            progress_frame,
            text="",
            font=('Microsoft YaHei UI', 9),
            foreground='#888'
        )
        self.progress_text.grid(row=2, column=0, sticky=tk.W, pady=(0, 8))

        # 统计信息区域 - 三格（成功/失败/跳过）。"总计"信息已合并到上方 status_label
        # 的 "X / Y" 后缀里，避免重复。方格改为无边框浅底色，降低视觉饱和度。
        stats_frame = ttk.Frame(progress_frame)
        stats_frame.grid(row=3, column=0, sticky=(tk.W, tk.E), pady=(4, 0))
        for col in range(3):
            stats_frame.grid_columnconfigure(col, weight=1)

        def _stat_cell(col, label_text, bg, number_fg, pad):
            cell = tk.Frame(stats_frame, bg=bg, relief='flat', borderwidth=0)
            cell.grid(row=0, column=col, sticky=(tk.W, tk.E), padx=pad)
            tk.Label(
                cell,
                text=label_text,
                font=('Microsoft YaHei UI', 9),
                bg=bg,
                fg='#666'
            ).pack(pady=(8, 2))
            number = tk.Label(
                cell,
                text="0",
                font=('Microsoft YaHei UI', 16, 'bold'),
                bg=bg,
                fg=number_fg
            )
            number.pack(pady=(2, 8))
            return number

        self.success_label = _stat_cell(0, "成功", '#EAF7EE', '#27ae60', (0, 4))
        self.failed_label = _stat_cell(1, "失败", '#FDECEC', '#e74c3c', 4)
        self.skipped_label = _stat_cell(2, "跳过", '#FFF4E2', '#f57c00', (4, 0))

        # 日志面板
        log_container = ttk.Frame(progress_frame)
        log_container.grid(row=4, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(8, 0))
        log_container.grid_columnconfigure(0, weight=1)
        log_container.grid_rowconfigure(0, weight=1)

        self.log_text = tk.Text(
            log_container,
            height=10,
            # Emoji 字形由 Segoe UI Emoji 提供；CHAR 换行可处理没有空格的长路径/英文。
            wrap=tk.CHAR,
            font=('Segoe UI Emoji', 9),
            bg='#1e1e1e',
            fg='#d4d4d4',
            insertbackground='#d4d4d4',
            selectbackground='#264f78',
            relief='flat',
            borderwidth=0,
            padx=8,
            pady=8,
            state='disabled'
        )
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        log_scrollbar = ttk.Scrollbar(log_container, orient='vertical', command=self.log_text.yview)
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.log_text.config(yscrollcommand=log_scrollbar.set)

        # 输出目录按钮区域
        output_frame = ttk.LabelFrame(main_frame, text="📂 输出目录", padding="12")
        output_frame.grid(row=5, column=0, pady=(0, 12), sticky=(tk.W, tk.E))
        output_frame.grid_columnconfigure(0, weight=1)
        output_frame.grid_columnconfigure(1, weight=1)

        # 显示当前输出根目录（灰色小字），用户能一眼看到产物会落在哪
        ttk.Label(
            output_frame,
            text=f"根目录：{self.output_path}",
            font=('Microsoft YaHei UI', 8),
            foreground='#888'
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, pady=(0, 6))

        ttk.Button(
            output_frame,
            text="📁 打开原始输出 (raw)",
            command=lambda: self.open_folder(self.raw_output_path)
        ).grid(row=1, column=0, sticky=(tk.W, tk.E), padx=(0, 6), pady=4)

        ttk.Button(
            output_frame,
            text="📁 打开提取结果 (extract)",
            command=lambda: self.open_folder(self.extract_output_path)
        ).grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(6, 0), pady=4)

        ttk.Button(
            output_frame,
            text="合并同名章节和图表到 Markdown",
            command=self.merge_all_sections_and_charts_to_markdown,
        ).grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(8, 0))

        # 底部信息
        footer_frame = ttk.Frame(main_frame)
        footer_frame.grid(row=6, column=0, sticky=(tk.W, tk.E), pady=(8, 0))

        ttk.Label(
            footer_frame,
            text=f"联系邮箱：{__contact_email__}",
            font=('Microsoft YaHei UI', 8),
            foreground='#999'
        ).pack(side=tk.LEFT)

    def _configured_llm_models(self) -> List[str]:
        """返回当前接口中允许在主界面选择的模型。"""
        if not self.llm_settings:
            return []
        return list(self.llm_settings.enabled_models)

    def _llm_provider_display_name(self) -> str:
        if not self.llm_settings:
            return "不可用"
        return "DeepSeek" if self.llm_settings.provider == "deepseek" else "自定义接口"

    def _refresh_llm_controls(self):
        """设置保存后刷新主界面的接口名称和模型下拉框。"""
        if not LLM_AVAILABLE:
            return
        self.llm_settings = load_llm_settings(self.base_path / ".env")
        models = self._configured_llm_models()
        active = self.llm_settings.active_model
        self.llm_model_combo.configure(
            values=models,
            state="readonly" if models else "disabled",
        )
        self.llm_model_var.set(active if active in models else (models[0] if models else ""))
        self.llm_provider_label.configure(
            text=f"接口：{self._llm_provider_display_name()}（在设置中管理）"
        )

    def _on_llm_model_selected(self, _event=None):
        """把主界面选中的自定义模型持久化为当前模型。"""
        if not self.llm_settings:
            return
        selected = self.llm_model_var.get().strip()
        if not selected:
            return
        if self.llm_settings.provider == "custom":
            self.llm_settings.custom_api_model = selected
        else:
            self.llm_settings.deepseek_model = selected
        try:
            save_llm_settings(self.llm_settings, self.base_path / ".env")
        except OSError as exc:
            self.log(f"⚠️  无法保存当前 LLM 模型: {exc}")

    def merge_all_sections_and_charts_to_markdown(self):
        """合并同名章节及各论文 Word 文件夹中的图表 Markdown。"""
        target = self.extract_output_path / "MergedSections"
        try:
            (
                outputs,
                total_section_articles,
                section_count,
                chart_count,
            ) = merge_sections_and_charts_to_markdown_files(
                self.extract_output_path,
                target,
            )
        except ValueError as exc:
            messagebox.showinfo("没有可合并内容", str(exc))
            return
        except Exception as exc:
            self.log(f"❌ 合并 Markdown 失败: {exc}")
            messagebox.showerror("合并失败", str(exc))
            return
        self.log(f"✅ 已生成 {len(outputs)} 个合并 Markdown")
        self.log(
            f"  - 同名章节：{section_count} 类，累计写入 {total_section_articles} 篇文章章节"
        )
        self.log(f"  - 图表汇总：合并 {chart_count} 个 Word 文件夹 Markdown")
        for output in outputs:
            self.log(f"  - {output}")
        messagebox.showinfo(
            "合并完成",
            (
                f"已生成 {len(outputs)} 个 Markdown 文件。\n"
                f"同名章节：{section_count} 类；图表来源：{chart_count} 个。\n"
                f"输出目录：{target}"
            ),
        )

    def check_input_folder(self):
        """检查用户选择的输入目录中的 PDF 文件。"""
        try:
            if not self.input_path.exists():
                self.input_path.mkdir(parents=True, exist_ok=True)
        except OSError as exc:
            self.file_count_label.config(text="PDF 文件数量: 0")
            self.log(f"[ERROR] 无法访问输入目录 {self.input_path}: {exc}")
            return

        try:
            pdf_files = sorted(self.input_path.glob("*.pdf"), key=lambda p: p.name.lower())
        except OSError as exc:
            self.file_count_label.config(text="PDF 文件数量: 0")
            self.log(f"[ERROR] 无法读取输入目录 {self.input_path}: {exc}")
            return
        count = len(pdf_files)
        self.file_count_label.config(text=f"PDF 文件数量: {count}")

        if count > 0:
            self.log(f"✓ 找到 {count} 个 PDF 文件")
            for pdf in pdf_files:
                self.log(f"  - {pdf.name}")
        else:
            self.log(f"⚠️  输入目录中没有 PDF 文件: {self.input_path}")

    LOG_MAX_LINES = 1000
    # Tkinter on Windows cannot reliably render color Emoji in every font.
    # Keep the original console messages, but use portable markers in the GUI.
    GUI_LOG_MARKERS = {
        "⚠️": "[WARN]",
        "⚠": "[WARN]",
        "✅": "[OK]",
        "❌": "[ERROR]",
        "🔍": "[SCAN]",
        "🤖": "[LLM]",
        "📈": "[STATS]",
        "📋": "[TEXT]",
        "🔄": "[RETRY]",
        "⏳": "[WAIT]",
        "💡": "[TIP]",
        "ℹ️": "[INFO]",
        "ℹ": "[INFO]",
        "⏭️": "[SKIP]",
        "⏭": "[SKIP]",
        "↳": "->",
    }

    def log(self, message: str):
        """添加日志消息到 GUI 和 UTF-8 日志文件（线程安全）。"""
        message = str(message)
        if sys.stdout is not None:
            print(message, flush=True)
        if self._log_handle is not None:
            try:
                with self._log_lock:
                    self._log_handle.write(message + '\n')
                    self._log_handle.flush()
            except (OSError, ValueError):
                pass
        self._post_ui(self._append_log, message)

    def _initialize_file_logging(self):
        """为本次 GUI 会话创建独立日志文件。"""
        try:
            log_directory = self.base_path / "logs"
            log_directory.mkdir(parents=True, exist_ok=True)
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            self.log_file_path = log_directory / f"PaperMiner_{timestamp}.log"
            self._log_handle = self.log_file_path.open(
                "a", encoding="utf-8", buffering=1
            )
            try:
                faulthandler.enable(file=self._log_handle, all_threads=True)
                self._fault_handler_enabled = True
            except (OSError, RuntimeError, ValueError):
                self._fault_handler_enabled = False
            atexit.register(self._close_log_file)
        except OSError as exc:
            self.log_file_path = None
            self._log_handle = None
            _STARTUP_MESSAGES.append(f"警告: 无法创建日志文件: {exc}")

    def _close_log_file(self):
        handle = self._log_handle
        if handle is None:
            return
        try:
            if self._fault_handler_enabled:
                try:
                    faulthandler.disable()
                except RuntimeError:
                    pass
                self._fault_handler_enabled = False
            with self._log_lock:
                handle.flush()
                handle.close()
        except (OSError, ValueError):
            pass
        finally:
            self._log_handle = None

    def _install_exception_hooks(self):
        """把 Tk 回调和后台线程的未捕获异常写入软件日志面板。"""
        def report_tk_exception(exc_type, exc_value, exc_traceback):
            details = ''.join(
                traceback.format_exception(exc_type, exc_value, exc_traceback)
            ).rstrip()
            self.log("[ERROR] GUI 回调发生未捕获异常")
            self.log(details)
            messagebox.showerror("PaperMiner 错误", str(exc_value), parent=self.root)

        def report_thread_exception(args):
            details = ''.join(
                traceback.format_exception(
                    args.exc_type, args.exc_value, args.exc_traceback
                )
            ).rstrip()
            self.log(f"[ERROR] 后台线程 {args.thread.name} 发生未捕获异常")
            self.log(details)

        self.root.report_callback_exception = report_tk_exception
        if hasattr(threading, "excepthook"):
            threading.excepthook = report_thread_exception

    def _append_log(self, message: str):
        """在主线程中向日志面板追加文本（仅由 root.after 调用）"""
        for marker, replacement in self.GUI_LOG_MARKERS.items():
            message = message.replace(marker, replacement)
        self.log_text.config(state='normal')
        self.log_text.insert(tk.END, message + '\n')

        # 裁剪超出上限的旧行
        line_count = int(self.log_text.index('end-1c').split('.')[0])
        if line_count > self.LOG_MAX_LINES:
            excess = line_count - self.LOG_MAX_LINES
            self.log_text.delete('1.0', f'{excess}.0')

        self.log_text.see(tk.END)
        self.log_text.config(state='disabled')

    def is_already_processed(self, pdf_name: str) -> bool:
        """检查PDF是否已经处理过（extract目录存在且包含文件）"""
        extract_dir = self.extract_output_path / pdf_name
        if not extract_dir.exists():
            return False
        try:
            for item in extract_dir.rglob('*'):
                if item.is_file():
                    return True
        except OSError:
            return False
        return False

    def open_folder(self, folder_path: Path):
        """打开文件夹"""
        if folder_path.exists():
            import os
            os.startfile(str(folder_path))
        else:
            messagebox.showwarning("警告", f"目录不存在：{folder_path}")

    def on_mode_change(self):
        """处理模式切换"""
        mode = self.process_mode_var.get()

        if mode == "extract_only":
            # 仅提取模式：禁用 GPU 和后端选项
            self.gpu_checkbox.config(state='disabled')
            self.use_gpu_var.set(False)
            self.backend_combo.config(state='disabled')
        else:
            # 完整处理模式：启用 GPU 和后端选项
            self.gpu_checkbox.config(state='normal')
            self.use_gpu_var.set(True)
            self.backend_combo.config(state='readonly')

    def start_processing(self):
        """开始处理"""
        mode = self.process_mode_var.get()

        # 检查是否至少选择了一项提取内容
        if not any([
            self.extract_text_var.get(),
            self.extract_formula_var.get(),
            self.extract_figures_var.get(),
            self.extract_tables_var.get(),
            self.extract_sections_var.get()
        ]):
            messagebox.showwarning(
                "未选择提取项",
                "请至少选择一项提取内容！"
            )
            return

        if mode == "full":
            # 完整处理模式：需要 PDF 文件
            pdf_files = sorted(self.input_path.glob("*.pdf"), key=lambda p: p.name.lower())
            if not pdf_files:
                messagebox.showwarning(
                    "没有文件",
                    f"所选输入目录中没有 PDF 文件：\n{self.input_path}\n\n"
                    "请添加文件或重新选择输入目录。"
                )
                return

            # 构建提取内容描述
            extract_items = []
            if self.extract_text_var.get():
                extract_items.append("文字")
            if self.extract_formula_var.get():
                extract_items.append("公式")
            if self.extract_figures_var.get():
                extract_items.append("图片")
            if self.extract_tables_var.get():
                extract_items.append("表格")
            if self.extract_sections_var.get():
                extract_items.append("论文章节")

            extract_desc = "、".join(extract_items)

            # 计算跳过数量
            if self.skip_processed_var.get():
                skip_count = sum(1 for f in pdf_files if self.is_already_processed(f.stem))
                actual_count = len(pdf_files) - skip_count
            else:
                skip_count = 0
                actual_count = len(pdf_files)

            # 确认开始处理
            if skip_count > 0:
                confirm_msg = (
                    f"共 {len(pdf_files)} 个 PDF 文件。\n\n"
                    f"将跳过: {skip_count} 个 (已有处理结果)\n"
                    f"实际处理: {actual_count} 个\n\n"
                    f"提取项目：{extract_desc}\n\n是否继续？"
                )
            else:
                confirm_msg = (
                    f"将完整处理 {len(pdf_files)} 个 PDF 文件。\n\n"
                    f"提取项目：{extract_desc}\n\n"
                    f"这将运行 MinerU 并提取内容。\n\n是否继续？"
                )

            if not messagebox.askyesno("确认处理", confirm_msg):
                return

            items_to_process = pdf_files

        else:  # extract_only
            # 仅提取模式：检查 raw 文件夹
            if not self.raw_output_path.exists():
                messagebox.showwarning(
                    "没有 raw 文件夹",
                    f"raw 文件夹不存在：{self.raw_output_path}\n\n"
                    f"请先运行完整处理模式生成 raw 文件夹。"
                )
                return

            # 查找所有 raw 子文件夹
            raw_folders = [d for d in self.raw_output_path.iterdir() if d.is_dir()]
            if not raw_folders:
                messagebox.showwarning(
                    "没有数据",
                    f"raw 文件夹中没有数据！\n\n"
                    f"请先运行完整处理模式。"
                )
                return

            # 构建提取内容描述
            extract_items = []
            if self.extract_text_var.get():
                extract_items.append("文字")
            if self.extract_formula_var.get():
                extract_items.append("公式")
            if self.extract_figures_var.get():
                extract_items.append("图片")
            if self.extract_tables_var.get():
                extract_items.append("表格")
            if self.extract_sections_var.get():
                extract_items.append("论文章节")

            extract_desc = "、".join(extract_items)

            # 计算跳过数量
            if self.skip_processed_var.get():
                skip_count = sum(1 for f in raw_folders if self.is_already_processed(f.name))
                actual_count = len(raw_folders) - skip_count
            else:
                skip_count = 0
                actual_count = len(raw_folders)

            # 确认开始提取
            if skip_count > 0:
                confirm_msg = (
                    f"共 {len(raw_folders)} 个 raw 文件夹。\n\n"
                    f"将跳过: {skip_count} 个 (已有提取结果)\n"
                    f"实际处理: {actual_count} 个\n\n"
                    f"提取项目：{extract_desc}\n\n是否继续？"
                )
            else:
                confirm_msg = (
                    f"将从 {len(raw_folders)} 个 raw 文件夹中提取内容。\n\n"
                    f"提取项目：{extract_desc}\n\n"
                    f"这将跳过 MinerU 处理，直接从已有的 raw 文件夹提取。\n\n是否继续？"
                )

            if not messagebox.askyesno("确认提取", confirm_msg):
                return

            items_to_process = raw_folders

        # 更新界面状态
        self._capture_run_options()
        self.is_processing = True
        self.start_button.config(state='disabled')
        self.stop_button.config(state='normal')

        # 重置计数器和统计显示
        self.current_pdf_index = 0
        self.total_pdfs = len(items_to_process)
        self.success_count = 0
        self.failed_count = 0
        self.skipped_count = 0

        # 更新统计显示
        self.update_stats()

        # 在后台线程中处理
        if mode == "full":
            threading.Thread(
                target=self.process_pdfs,
                args=(items_to_process,),
                daemon=True,
                name="PaperMiner-PDF-Batch",
            ).start()
        else:
            threading.Thread(
                target=self.extract_from_raw,
                args=(items_to_process,),
                daemon=True,
                name="PaperMiner-Raw-Extract",
            ).start()

    def stop_processing(self):
        """停止处理"""
        self.is_processing = False
        self.log("\n⚠️  用户请求停止处理...")
        self._terminate_active_mineru()

    def process_pdfs(self, pdf_files: List[Path]):
        """处理所有 PDF 文件"""
        try:
            self.log("=" * 60)
            self.log("PaperMiner - 批量 PDF 处理开始")
            self.log("=" * 60)
            self.log(f"总文件数: {len(pdf_files)}")
            self.log(f"输出目录: {self.output_path}")
            self.log("")

            # 创建输出目录
            self.raw_output_path.mkdir(parents=True, exist_ok=True)
            self.extract_output_path.mkdir(parents=True, exist_ok=True)

            # 检查 GPU 状态
            self.precheck_mineru_environment()
            self.check_gpu_status()

            # 处理每个 PDF
            for i, pdf_file in enumerate(pdf_files):
                if not self.is_processing:
                    self.log("\n❌ 处理已停止")
                    break

                self.current_pdf_index = i + 1
                self._post_ui(self.update_progress)

                # 检查是否跳过已处理的文件
                if self._run_option("skip_processed", True) and self.is_already_processed(pdf_file.stem):
                    self.skipped_count += 1
                    self.log(f"\n⏭ 跳过已处理: {pdf_file.name}")
                    self._post_ui(self.update_stats)
                    continue

                self.log("\n" + "=" * 60)
                self.log(f"[{i+1}/{len(pdf_files)}] 处理: {pdf_file.name}")
                self.log("=" * 60)

                # 步骤 1: 使用 MinerU 处理 PDF
                raw_dir = self.run_mineru(pdf_file)

                if raw_dir is not None:
                    # 步骤 2: 提取和整理结果
                    extract_ok = self.extract_and_organize(pdf_file.stem, raw_dir=raw_dir)
                    if extract_ok:
                        self.success_count += 1
                        self.log(f"✅ 完成: {pdf_file.name}")
                    else:
                        self.failed_count += 1
                        self.log(f"❌ 提取失败: {pdf_file.name}")
                else:
                    self.failed_count += 1
                    self.log(f"❌ 失败: {pdf_file.name}")
                    if self.last_mineru_issue_code in {
                        "model_snapshot_missing",
                        "hf_network_error",
                        "mineru_config_missing",
                        "mineru_missing",
                        "unsupported_mineru_version",
                        "mineru_worker_missing",
                    }:
                        self.log("⚠️  检测到环境级错误，停止后续文件处理。请先修复环境后重试。")
                        self._release_batch_memory(pdf_file.name)
                        break

                # 更新统计显示
                self._post_ui(self.update_stats)
                self._release_batch_memory(pdf_file.name)

            # 处理完成
            self.log("\n" + "=" * 60)
            self.log("处理完成!")
            self.log("=" * 60)
            self.log(f"成功: {self.success_count} 个")
            self.log(f"失败: {self.failed_count} 个")
            if self.skipped_count > 0:
                self.log(f"跳过: {self.skipped_count} 个")
            self.log(f"总计: {len(pdf_files)} 个")

            self._post_ui(self.processing_complete)

        except Exception as e:
            self.log(f"\n❌ 处理过程中发生错误: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            self._post_ui(self.processing_complete)

    def update_progress(self):
        """更新进度。status_label 承载"状态 · X / Y"总进度；progress_text 留给
        具体的当前文件名（由调用方按需写入），避免两者信息打架。"""
        if self.total_pdfs > 0:
            progress = (self.current_pdf_index / self.total_pdfs) * 100
            self.progress_var.set(progress)
            self.status_label.config(
                text=f"处理中 · {self.current_pdf_index} / {self.total_pdfs}",
                foreground=self.accent_color
            )

    def update_stats(self):
        """更新统计信息显示（总计已合并到 status_label，不再单独展示）"""
        self.success_label.config(text=str(self.success_count))
        self.failed_label.config(text=str(self.failed_count))
        self.skipped_label.config(text=str(self.skipped_count))

    @staticmethod
    def _process_memory_snapshot():
        """返回当前进程的工作集/私有内存（MiB）；诊断失败时返回 None。"""
        if sys.platform != "win32":
            return None
        try:
            import ctypes
            from ctypes import wintypes

            class ProcessMemoryCountersEx(ctypes.Structure):
                _fields_ = [
                    ("cb", wintypes.DWORD),
                    ("PageFaultCount", wintypes.DWORD),
                    ("PeakWorkingSetSize", ctypes.c_size_t),
                    ("WorkingSetSize", ctypes.c_size_t),
                    ("QuotaPeakPagedPoolUsage", ctypes.c_size_t),
                    ("QuotaPagedPoolUsage", ctypes.c_size_t),
                    ("QuotaPeakNonPagedPoolUsage", ctypes.c_size_t),
                    ("QuotaNonPagedPoolUsage", ctypes.c_size_t),
                    ("PagefileUsage", ctypes.c_size_t),
                    ("PeakPagefileUsage", ctypes.c_size_t),
                    ("PrivateUsage", ctypes.c_size_t),
                ]

            counters = ProcessMemoryCountersEx()
            counters.cb = ctypes.sizeof(counters)
            kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
            psapi = ctypes.WinDLL("psapi", use_last_error=True)
            kernel32.GetCurrentProcess.restype = wintypes.HANDLE
            get_process_memory_info = psapi.GetProcessMemoryInfo
            get_process_memory_info.argtypes = [
                wintypes.HANDLE,
                ctypes.POINTER(ProcessMemoryCountersEx),
                wintypes.DWORD,
            ]
            get_process_memory_info.restype = wintypes.BOOL
            ok = get_process_memory_info(
                kernel32.GetCurrentProcess(),
                ctypes.byref(counters),
                counters.cb,
            )
            if not ok:
                return None
            divisor = 1024 * 1024
            return {
                "working_set": counters.WorkingSetSize / divisor,
                "private": counters.PrivateUsage / divisor,
            }
        except Exception:
            return None

    def _release_batch_memory(self, context: str):
        """每篇结束后释放临时对象和 CUDA 缓存，并记录内存变化。"""
        before = self._process_memory_snapshot()
        gpu_before = None
        gpu_after = None
        # 先打破/回收 Python 引用，再让 CUDA caching allocator 归还空闲块；
        # 顺序反过来会遗漏刚刚由 gc 释放的 tensor。
        collected = gc.collect()
        try:
            torch_module = sys.modules.get("torch")
            if torch_module is not None and torch_module.cuda.is_available():
                gpu_before = (
                    torch_module.cuda.memory_allocated() / (1024 * 1024),
                    torch_module.cuda.memory_reserved() / (1024 * 1024),
                )
                torch_module.cuda.empty_cache()
                if hasattr(torch_module.cuda, "ipc_collect"):
                    torch_module.cuda.ipc_collect()
                gpu_after = (
                    torch_module.cuda.memory_allocated() / (1024 * 1024),
                    torch_module.cuda.memory_reserved() / (1024 * 1024),
                )
        except Exception as exc:
            self.log(f"[WARN] CUDA 缓存清理失败（不影响后续处理）: {exc}")

        after = self._process_memory_snapshot()
        parts = [f"内存回收 [{context}]: Python 对象 {collected}"]
        if before and after:
            parts.append(
                "工作集 "
                f"{before['working_set']:.0f}->{after['working_set']:.0f} MiB，"
                f"私有内存 {before['private']:.0f}->{after['private']:.0f} MiB"
            )
        if gpu_before and gpu_after:
            parts.append(
                "GPU 已分配/保留 "
                f"{gpu_before[0]:.0f}/{gpu_before[1]:.0f}"
                f"->{gpu_after[0]:.0f}/{gpu_after[1]:.0f} MiB"
            )
        self.log("；".join(parts))

    def check_gpu_status(self):
        """检查 GPU 状态"""
        try:
            self.log("=== GPU 诊断 ===")

            result = subprocess.run([
                sys.executable, "-c",
                "import torch; "
                "print('PyTorch:', torch.__version__); "
                "print('CUDA:', torch.cuda.is_available()); "
                "print('GPU:', torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'N/A')"
            ], capture_output=True, text=True, encoding='utf-8', errors='replace', timeout=30)

            if result.returncode == 0:
                for line in result.stdout.strip().split('\n'):
                    self.log(line)

                # 检查 CUDA 是否可用
                cuda_check = subprocess.run([
                    sys.executable, "-c",
                    "import torch; exit(0 if torch.cuda.is_available() else 1)"
                ], capture_output=True, timeout=10)

                if cuda_check.returncode != 0:
                    self.log("⚠️  CUDA 不可用，将使用 CPU 模式")
                    self._active_run_options["use_gpu"] = False
                    self._post_ui(self.use_gpu_var.set, False)
                else:
                    self.log("✅ GPU 加速已启用")
            else:
                self.log("⚠️  GPU 检查失败")

            self.log("=" * 60)
            self.log("")

        except Exception as e:
            self.log(f"⚠️  GPU 状态检查失败: {str(e)}")
            self.log("")

    def run_mineru(self, pdf_file: Path) -> 'Path | None':
        """在隔离子进程中调用 MinerU，成功返回 raw_dir，失败返回 None。

        pipeline 会加载 CUDA、PyTorch、ONNX Runtime 等原生库。原生访问冲突或
        资源耗尽不会产生可捕获的 Python 异常；若直接在 GUI 进程内调用，整个
        PaperMiner 会无 traceback 地退出。每篇 PDF 使用独立进程，既保住主界面
        与日志，也确保 Windows 在每篇结束后回收全部进程级 CPU/GPU 资源。
        """
        output_tail: List[str] = []
        process = None
        try:
            self.last_mineru_issue_code = None
            self.log("步骤 1: 使用 MinerU 处理 PDF (稳定隔离进程)...")

            device = "cuda" if self._run_option("use_gpu", True) else "cpu"
            backend = self._run_option("backend", "pipeline")
            model_source = os.environ.get("MINERU_MODEL_SOURCE", "modelscope")
            worker_script = Path(__file__).resolve().with_name("mineru_worker.py")
            if not worker_script.is_file():
                self.last_mineru_issue_code = "mineru_worker_missing"
                self.log(f"❌ MinerU 隔离组件缺失: {worker_script}")
                self.log("  请从 PaperMiner 安装程序执行重装。")
                return None

            runtime_python = Path(sys.executable).resolve()
            if runtime_python.name.lower() == "pythonw.exe":
                console_python = runtime_python.with_name("python.exe")
                if console_python.is_file():
                    runtime_python = console_python

            command = (
                str(runtime_python),
                "-u",
                str(worker_script),
                "--input",
                str(pdf_file),
                "--output",
                str(self.raw_output_path),
                "--device",
                device,
                "--backend",
                backend,
                "--model-source",
                model_source,
            )
            worker_env = os.environ.copy()
            worker_env["MINERU_DEVICE_MODE"] = device
            worker_env["MINERU_MODEL_SOURCE"] = model_source
            worker_env["PYTHONIOENCODING"] = "utf-8"
            worker_env["PYTHONFAULTHANDLER"] = "1"
            worker_env["PYTHONNOUSERSITE"] = "1"

            creation_flags = 0
            if sys.platform == "win32":
                creation_flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)

            process = subprocess.Popen(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                encoding="utf-8",
                errors="replace",
                bufsize=1,
                env=worker_env,
                creationflags=creation_flags,
            )
            self._mineru_process = process
            if process.stdout is not None:
                for raw_line in process.stdout:
                    line = raw_line.rstrip("\r\n")
                    output_tail.append(line)
                    self.log(f"  {line}" if line else "")
                    if len(output_tail) > 500:
                        output_tail.pop(0)

            exit_code = process.wait()
            self._mineru_process = None
            if exit_code != 0:
                unsigned_code = exit_code & 0xFFFFFFFF
                code_text = f"{exit_code} (0x{unsigned_code:08X})"
                self.log("")
                self.log(f"❌ MinerU 隔离进程异常退出，代码: {code_text}")

                if exit_code == 11:
                    self.last_mineru_issue_code = "mineru_missing"
                elif exit_code == 12:
                    self.last_mineru_issue_code = "unsupported_mineru_version"
                elif unsigned_code >= 0x80000000:
                    self.last_mineru_issue_code = "mineru_native_crash"
                    native_hints = {
                        0xC0000005: "原生访问冲突（常见于 CUDA/显卡驱动/原生推理库）",
                        0xC0000017: "系统无法分配所需内存",
                        0xC000009A: "系统资源不足",
                        0xC00000FD: "原生线程堆栈溢出",
                        0xC0000374: "原生堆损坏",
                        0xC0000409: "原生库快速失败/安全检查失败",
                    }
                    hint = native_hints.get(
                        unsigned_code,
                        "CUDA/PyTorch/ONNX 等原生组件异常",
                    )
                    self.log(f"  判定: {hint}")
                    self.log("  主程序已被隔离保护；该文献记为失败，其余队列可继续。")
                else:
                    self.last_mineru_issue_code = "mineru_worker_failed"

                diagnosis = self.diagnose_mineru_output(output_tail)
                if diagnosis:
                    self.log_mineru_diagnosis(diagnosis, output_tail)
                return None

            pdf_name = pdf_file.stem
            self.log("")
            self.log("✓ MinerU 处理完成")
            raw_dir = self.find_raw_output_dir(pdf_name)
            if raw_dir and self.validate_mineru_output(raw_dir):
                self.log("✓ MinerU 输出验证通过")
                return raw_dir

            self.log("❌ MinerU 执行成功但未生成有效输出文件")
            diagnosis = self.diagnose_mineru_output(output_tail)
            if diagnosis:
                self.last_mineru_issue_code = diagnosis.get("code")
            self.log_mineru_diagnosis(diagnosis, output_tail)
            return None

        except Exception as e:
            self.log("")
            self.log(f"❌ MinerU 处理失败: {e}")
            tb = traceback.format_exc()
            self.log(tb)
            output_tail.append(str(e))
            output_tail.extend(tb.splitlines())
            diagnosis = self.diagnose_mineru_output(output_tail)
            if diagnosis:
                self.last_mineru_issue_code = diagnosis.get("code")
                self.log_mineru_diagnosis(diagnosis, output_tail)
            return None
        finally:
            self._mineru_process = None
            if process is not None and process.stdout is not None:
                try:
                    process.stdout.close()
                except (OSError, ValueError):
                    pass

    def precheck_mineru_environment(self):
        """预检 MinerU 模型配置，帮助用户在运行前发现常见环境问题。"""
        try:
            offline_flag = os.environ.get("HF_HUB_OFFLINE", "").strip().lower()
            if offline_flag in {"1", "true", "yes", "on"}:
                self.log("⚠️  [预检] 检测到 HF_HUB_OFFLINE=1，仅会使用本地缓存。")
                self.log("    若后续出现 LocalEntryNotFoundError，请先下载模型。")

            config_path = Path.home() / "mineru.json"
            if not config_path.exists():
                self.log("ℹ️  [预检] 未找到 mineru.json（不一定是问题）。")
                self.log("    若后续处理能成功，此提示可忽略；若失败请执行模型下载命令。")
                self.log("    建议命令: mineru-models-download --source modelscope --model_type pipeline")
                return

            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
            except Exception as e:
                self.log(f"⚠️  [预检] 读取 mineru.json 失败: {e}")
                self.log("    若后续处理失败，请重新执行模型下载命令以重建配置。")
                return

            model_dir = None
            if isinstance(config, dict):
                models_dir = config.get("models-dir")
                if isinstance(models_dir, dict):
                    model_dir = models_dir.get("pipeline")

            if not model_dir:
                self.log("⚠️  [预检] mineru.json 中缺少 models-dir.pipeline。")
                self.log("    若后续处理失败，请执行: mineru-models-download --source modelscope --model_type pipeline")
                self.last_mineru_issue_code = "mineru_config_missing"
                return

            model_dir_path = Path(model_dir).expanduser()
            if not model_dir_path.exists():
                self.log(f"⚠️  [预检] 模型目录不存在: {model_dir_path}")
                self.log("    若后续处理失败，请重新下载模型。")
                self.last_mineru_issue_code = "mineru_config_missing"
                return

            try:
                has_any_file = any(model_dir_path.rglob("*"))
            except OSError:
                has_any_file = False

            if not has_any_file:
                self.log(f"⚠️  [预检] 模型目录为空: {model_dir_path}")
                self.log("    若后续处理失败，请重新下载模型。")
                self.last_mineru_issue_code = "mineru_config_missing"

        except Exception as e:
            self.log(f"⚠️  [预检] MinerU 环境检查失败: {e}")

    def diagnose_mineru_output(self, output_lines: List[str]):
        """根据 MinerU 输出诊断常见失败原因。"""
        if not output_lines:
            return None

        text = "\n".join(output_lines).lower()

        if (
            "localentrynotfounderror" in text
            or "cannot find the appropriate snapshot folder" in text
            or "trying to locate the files on the hub" in text
        ):
            return {
                "code": "model_snapshot_missing",
                "title": "未找到本地模型快照（HuggingFace 缓存缺失或不可访问）",
                "tips": [
                    "MinerU 3.0+ 首次运行会自动下载模型，需要网络访问模型源",
                    "国内用户请确认环境变量: set MINERU_MODEL_SOURCE=modelscope （运行程序.bat 已自动设置）",
                    "也可手动下载: conda activate MinerU && mineru-models-download --source modelscope --model_type pipeline",
                    "若网络受限，在可联网机器下载模型后复制缓存目录到本机",
                ],
            }

        if (
            ("connectionerror" in text or "maxretryerror" in text or "proxyerror" in text or "timeout" in text)
            and ("huggingface" in text or "hf.co" in text or "hub" in text)
        ):
            return {
                "code": "hf_network_error",
                "title": "模型下载网络失败（HuggingFace 连接异常）",
                "tips": [
                    "国内用户请确认环境变量: set MINERU_MODEL_SOURCE=modelscope （运行程序.bat 已自动设置）",
                    "也可手动下载: conda activate MinerU && mineru-models-download --source modelscope --model_type pipeline",
                    "检查代理/防火墙设置，确保当前环境可访问模型源",
                ],
            }

        if (
            "cuda out of memory" in text
            or "outofmemoryerror" in text
            or ("cuda" in text and "allocate" in text and "memory" in text)
        ):
            return {
                "code": "cuda_oom",
                "title": "GPU 显存不足（CUDA Out of Memory）",
                "tips": [
                    "取消界面上'使用 GPU 加速'改为 CPU 模式再跑",
                    "或显式设置 MINERU_VIRTUAL_VRAM_SIZE=<可用显存GB>，按实际显存调小",
                    "批处理时一次处理一篇 PDF（当前已是单文件模式，若仍 OOM 请先降显存档位）",
                ],
            }

        return None

    def log_mineru_diagnosis(self, diagnosis, output_lines: List[str]):
        """输出诊断信息，帮助用户快速定位失败原因。"""
        keywords = ["error", "exception", "traceback", "failed", "notfound", "timeout"]
        recent_errors = [line for line in output_lines if any(k in line.lower() for k in keywords)][-8:]

        if recent_errors:
            self.log("  [诊断] 关键错误输出:")
            for line in recent_errors:
                self.log(f"    {line}")

        if not diagnosis:
            return

        self.log(f"  [诊断] {diagnosis['title']}")
        tips = diagnosis.get("tips", [])
        if tips:
            self.log("  [建议] 可按以下步骤处理:")
            for tip in tips:
                self.log(f"    - {tip}")

    def find_raw_output_dir(self, pdf_name: str) -> 'Path | None':
        """
        多策略查找 MinerU 输出目录。
        MinerU 可能因长文件名/特殊字符导致目录名与 pdf_name 不完全一致。
        """
        # 策略1：精确匹配
        exact = self.raw_output_path / pdf_name / "auto"
        if exact.exists():
            return exact

        # 策略2：前缀 glob 匹配（取前50字符，处理长文件名截断）
        prefix = pdf_name[:50].rstrip()
        try:
            candidates = [
                d / "auto" for d in self.raw_output_path.iterdir()
                if d.is_dir() and d.name.startswith(prefix) and (d / "auto").exists()
            ]
        except OSError:
            candidates = []

        if len(candidates) == 1:
            self.log(f"  [路径] 前缀匹配: {candidates[0]}")
            return candidates[0]
        elif len(candidates) > 1:
            best = max(candidates, key=lambda p: p.stat().st_mtime)
            self.log(f"  [路径] 前缀匹配到多个目录，使用最新的: {best}")
            return best

        # 策略3：查找最近10分钟内创建的 auto 目录
        now = time.time()
        recent = []
        try:
            for d in self.raw_output_path.iterdir():
                auto = d / "auto"
                if auto.exists() and (now - auto.stat().st_mtime) < 600:
                    recent.append(auto)
        except OSError:
            pass

        if len(recent) == 1:
            self.log(f"  [路径] 最近创建匹配: {recent[0]}")
            return recent[0]

        self.log(f"  [路径] ❌ 未找到匹配的输出目录 (pdf_name={pdf_name})")
        return None

    def validate_mineru_output(self, raw_dir: Path) -> bool:
        """验证 MinerU 输出目录中存在可用文件"""
        md_files = list(raw_dir.glob("*.md"))
        json_files = list(raw_dir.glob("*_content_list.json"))

        if not md_files and not json_files:
            self.log("  [验证] ❌ MinerU 输出目录为空 - 无 .md 或 _content_list.json")
            return False

        if md_files:
            self.log(f"  [验证] ✓ Markdown: {md_files[0].name}")
        if json_files:
            self.log(f"  [验证] ✓ content_list: {json_files[0].name}")
        return True

    def find_file_by_glob(self, directory: Path, exact_name: str, glob_pattern: str) -> 'Path | None':
        """先精确匹配文件名，再用 glob 兜底"""
        exact = directory / exact_name
        if exact.exists():
            return exact
        matches = list(directory.glob(glob_pattern))
        if matches:
            return matches[0]
        return None

    def extract_from_raw(self, raw_folders: List[Path]):
        """从已有的 raw 文件夹中提取内容"""
        try:
            self.log("=" * 60)
            self.log("仅提取模式 - 从 raw 文件夹提取内容")
            self.log("=" * 60)
            self.log(f"总文件夹数: {len(raw_folders)}")
            self.log(f"输出目录: {self.extract_output_path}")
            self.log("=" * 60)
            self.log("")

            for i, raw_folder in enumerate(raw_folders, 1):
                if not self.is_processing:
                    self.log("\n⚠️  处理已停止")
                    break

                self.current_pdf_index = i
                pdf_name = raw_folder.name

                # 更新进度
                self._post_ui(self.update_progress)

                # 检查是否跳过已处理的文件
                if self._run_option("skip_processed", True) and self.is_already_processed(pdf_name):
                    self.skipped_count += 1
                    self.log(f"⏭ 跳过已处理: {pdf_name}")
                    self._post_ui(self.update_stats)
                    continue

                self.log("=" * 60)
                self.log(f"[{i}/{len(raw_folders)}] 提取: {pdf_name}")
                self.log("=" * 60)

                # 更新进度：status_label 由 update_progress 写成 "处理中 · X/Y"，
                # 文件名单独落在 progress_text（副行灰字），两者不再打架。
                self._post_ui(self.update_progress)
                self._post_ui(self.progress_text.config, text=f"正在提取：{pdf_name}")

                try:
                    # 提取内容（使用 PDF 名称作为参数）
                    extract_ok = self.extract_and_organize(pdf_name)

                    if extract_ok:
                        self.success_count += 1
                        self.log(f"✅ 完成: {pdf_name}")
                    else:
                        self.failed_count += 1
                        self.log(f"❌ 提取失败: {pdf_name}")
                    self.log("")

                except Exception as e:
                    self.failed_count += 1
                    self.log(f"❌ 失败: {pdf_name}")
                    self.log(f"   错误: {str(e)}")
                    import traceback
                    self.log(traceback.format_exc())
                    self.log("")

                # 更新统计显示
                self._post_ui(self.update_stats)
                self._release_batch_memory(pdf_name)

            # 处理完成
            self.log("=" * 60)
            self.log("提取完成!")
            self.log("=" * 60)
            self.log(f"成功: {self.success_count} 个")
            self.log(f"失败: {self.failed_count} 个")
            if self.skipped_count > 0:
                self.log(f"跳过: {self.skipped_count} 个")
            self.log(f"总计: {len(raw_folders)} 个")

            self._post_ui(self.processing_complete)

        except Exception as e:
            self.log(f"❌ 提取过程出错: {str(e)}")
            self._post_ui(self.processing_complete)

    def extract_and_organize(self, pdf_name: str, raw_dir: 'Path | None' = None) -> bool:
        """提取和整理处理结果到extract文件夹。返回 True 表示至少部分提取成功。"""
        try:
            self.log("步骤 2: 提取和整理结果...")

            # 查找 MinerU 的输出目录
            if raw_dir is None:
                raw_dir = self.find_raw_output_dir(pdf_name)

            if raw_dir is None or not raw_dir.exists():
                self.log(f"⚠️  未找到输出目录 (pdf_name={pdf_name})")
                return False

            raw_pdf_dir = raw_dir

            # 从实际目录名推导 pdf_name（可能与传入的不同）
            actual_pdf_name = raw_dir.parent.name
            if actual_pdf_name != pdf_name:
                self.log(f"  [路径] 实际目录名与传入名称不同: {actual_pdf_name}")

            # 创建提取目录（在extract下为每个PDF创建子文件夹）
            extract_pdf_dir = self.extract_output_path / pdf_name
            extract_pdf_dir.mkdir(parents=True, exist_ok=True)

            any_success = False

            # 提取文字 (Markdown) - 保存到extract/pdf_name/pdf_name.md
            if self._run_option("extract_text", True):
                if self.extract_text(raw_pdf_dir, extract_pdf_dir, actual_pdf_name):
                    any_success = True

            # 提取公式 - 保存到extract/pdf_name/Formula/
            if self._run_option("extract_formula", True):
                if self.extract_formulas(raw_pdf_dir, extract_pdf_dir, actual_pdf_name):
                    any_success = True

            # 提取图片 - 保存到extract/pdf_name/Figure/
            if self._run_option("extract_figures", True):
                if self.extract_figures(raw_pdf_dir, extract_pdf_dir, actual_pdf_name):
                    any_success = True

            # 提取表格 - 保存到extract/pdf_name/Tables/
            if self._run_option("extract_tables", True):
                if self.extract_tables(raw_pdf_dir, extract_pdf_dir, actual_pdf_name):
                    any_success = True

            # 先创建 Word 文件夹（不依赖 LLM），避免 LLM 长耗时或失败时丢失 Word/docx 产出
            if self._run_option("extract_figures", True) or self._run_option("extract_tables", True):
                self.create_word_folder(raw_pdf_dir, extract_pdf_dir, actual_pdf_name)

            # 最后跑 LLM 章节提取（可能很慢，放最后即使超时/失败也不影响上面的产出）
            if self._run_option("extract_sections", True):
                self.extract_sections_with_llm(raw_pdf_dir, extract_pdf_dir, actual_pdf_name)

            if any_success:
                self.log("✓ 提取和整理完成")
            else:
                self.log("⚠️  所有提取步骤均未产生输出")

            return any_success

        except Exception as e:
            self.log(f"❌ 提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def processing_complete(self):
        """处理完成"""
        self.is_processing = False
        self.start_button.config(state='normal')
        self.stop_button.config(state='disabled')

        # 更新状态标签（附 X / Y 进度计数）；清空副行的"正在提取…"文件名提示
        done = self.success_count + self.failed_count + self.skipped_count
        total = self.total_pdfs if self.total_pdfs else done
        progress_suffix = f" · {done} / {total}" if total else ""
        if self.failed_count == 0:
            self.status_label.config(
                text=f"处理完成{progress_suffix}",
                foreground='#27ae60'
            )
        elif self.success_count == 0:
            self.status_label.config(
                text=f"处理失败{progress_suffix}",
                foreground='#e74c3c'
            )
        else:
            self.status_label.config(
                text=f"处理完成（部分失败）{progress_suffix}",
                foreground='#f39c12'
            )
        self.progress_text.config(text="")

        # 更新统计显示
        self.update_stats()

        # 显示完成消息
        summary_parts = [
            f"成功: {self.success_count} 个",
            f"失败: {self.failed_count} 个",
        ]
        if self.skipped_count > 0:
            summary_parts.append(f"跳过: {self.skipped_count} 个")
        summary_parts.append(f"总计: {self.total_pdfs} 个")

        messagebox.showinfo(
            "处理完成",
            "批量处理完成！\n\n"
            + "\n".join(summary_parts)
            + f"\n\n结果已保存到：\n{self.extract_output_path}"
        )

    def extract_text(self, raw_dir: Path, extract_dir: Path, pdf_name: str) -> bool:
        """提取文字 (Markdown)，并修复图片引用路径。

        策略：和 extract_figures 一致的阅读顺序分组，把 image+chart 子面板统一重定向到
        Figure/Fig.N.jpg（extract_figures 会重建该整图）。Table/Equation 路径同样重写。
        """
        try:
            self.log("  - 提取文字...")

            md_file = self.find_file_by_glob(raw_dir, f"{pdf_name}.md", "*.md")
            if md_file is None:
                self.log("    ⚠️  未找到 Markdown 文件")
                return False

            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()

            content_list_file = self.find_file_by_glob(raw_dir, f"{pdf_name}_content_list.json", "*_content_list.json")
            image_mapping = {}

            if content_list_file:
                with open(content_list_file, 'r', encoding='utf-8') as f:
                    content_list = json.load(f)

                # ---- 图片 (image+chart)：按阅读顺序分组，每组并入同一个 Fig.N.jpg ----
                fig_caption_pat = re.compile(
                    r'\b(Fig\.?|Figure|Scheme|图|示意图|插图)\s*(\d+)', re.IGNORECASE
                )
                page_items = {}
                for ci, item in enumerate(content_list):
                    if not isinstance(item, dict) or item.get('type') not in ('image', 'chart'):
                        continue
                    p = item.get('page_idx')
                    if p is None:
                        continue
                    caps = (
                        item.get('image_caption')
                        or item.get('chart_caption')
                        or item.get('img_caption')
                        or []
                    )
                    cap = ' '.join(caps).strip() if caps else ''
                    m = fig_caption_pat.search(cap) if cap else None
                    fig_num = m.group(2) if m else None
                    page_items.setdefault(p, []).append((ci, item, fig_num))

                for p, items in page_items.items():
                    pending = []
                    groups = []  # [(fig_num, [content_list_items])]
                    for ci, item, fn in items:
                        if fn:
                            groups.append((fn, pending + [item]))
                            pending = []
                        else:
                            pending.append(item)
                    if pending and groups:
                        lf, li = groups[-1]
                        groups[-1] = (lf, li + pending)
                    for fig_num, its in groups:
                        new_path = f"Figure/Fig.{fig_num}.jpg"
                        for it in its:
                            ip = it.get('img_path')
                            if ip:
                                image_mapping[ip] = new_path

                # ---- 表格 ----
                table_index = 1
                for item in content_list:
                    if not isinstance(item, dict) or item.get('type') != 'table':
                        continue
                    if 'img_path' not in item:
                        continue
                    caption = ' '.join(item.get('table_caption', []) or [])
                    m = re.search(r'Table\s*(\d+)', caption, re.IGNORECASE)
                    num = m.group(1) if m else str(table_index)
                    image_mapping[item['img_path']] = f"Tables/Table_{num}.jpg"
                    table_index += 1

                # ---- 公式 ----
                formula_index = 1
                for item in content_list:
                    if not isinstance(item, dict) or item.get('type') != 'equation':
                        continue
                    if 'img_path' not in item:
                        continue
                    image_mapping[item['img_path']] = f"Formula/formula_{formula_index}.jpg"
                    formula_index += 1

            for old_path, new_path in image_mapping.items():
                content = content.replace(f"]({old_path})", f"]({new_path})")

            output_md = extract_dir / f"{pdf_name}.md"
            with open(output_md, 'w', encoding='utf-8') as f:
                f.write(content)

            self.log(f"    ✓ 已保存: {output_md.name}")
            self.log(f"    ✓ 修复了 {len(image_mapping)} 个图片引用")
            return True

        except Exception as e:
            self.log(f"    ❌ 文字提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def extract_formulas(self, raw_dir: Path, extract_dir: Path, pdf_name: str) -> bool:
        """提取公式（从 content_list.json 和 Markdown）"""
        try:
            self.log("  - 提取公式...")

            # 读取 content_list.json
            content_list_file = self.find_file_by_glob(raw_dir, f"{pdf_name}_content_list.json", "*_content_list.json")
            formula_images = []

            if content_list_file:
                with open(content_list_file, 'r', encoding='utf-8') as f:
                    content_list = json.load(f)

                # 提取所有公式图片
                for item in content_list:
                    if item.get('type') == 'equation':
                        if 'img_path' in item:
                            formula_images.append({
                                'img_path': item['img_path'],
                                'latex': item.get('latex_text', '')
                            })

            # 同时从 Markdown 文件中提取文本公式
            md_file = self.find_file_by_glob(raw_dir, f"{pdf_name}.md", "*.md")
            text_formulas = []
            if md_file:
                with open(md_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                # 提取公式块 ($$...$$)
                text_formulas = re.findall(r'\$\$(.*?)\$\$', content, re.DOTALL)

            if not formula_images and not text_formulas:
                self.log("    ⚠️  未找到公式")
                return False

            # 创建公式文件夹
            formula_dir = extract_dir / "Formula"
            formula_dir.mkdir(exist_ok=True)

            # 保存公式图片
            if formula_images:
                for i, formula_item in enumerate(formula_images, 1):
                    img_path = raw_dir / formula_item['img_path']
                    if img_path.exists():
                        dest_path = formula_dir / f"formula_{i}{img_path.suffix}"
                        shutil.copy2(img_path, dest_path)

            # 保存文本公式到 Markdown 文件
            if text_formulas:
                formula_md = formula_dir / f"{pdf_name}_formula.md"
                with open(formula_md, 'w', encoding='utf-8') as f:
                    f.write(f"# {pdf_name} - 公式\n\n")
                    for i, formula in enumerate(text_formulas, 1):
                        f.write(f"## 公式 {i}\n\n")
                        f.write(f"$$\n{formula.strip()}\n$$\n\n")

            total_formulas = len(formula_images) + len(text_formulas)
            self.log(f"    ✓ 提取 {total_formulas} 个公式 (图片: {len(formula_images)}, 文本: {len(text_formulas)})")
            return True

        except Exception as e:
            self.log(f"    ❌ 公式提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def extract_figures(self, raw_dir: Path, extract_dir: Path, pdf_name: str) -> bool:
        """提取图片：为 mineru 3.x 拆碎的子面板重建完整 Figure。

        策略（按优先级）：
        1. 若 content_list 里出现带 'Fig/Figure/图 N' 标题的 image/chart 条目，且 `_origin.pdf`
           可用：对该页所有 image+chart 块的 bbox 求并集，从原 PDF 重渲染整块区域，
           保存为 `Figure/Fig.N.jpg`。这能在 mineru 3.x 下还原 2.5.3 的完整 Figure 行为。
        2. 若上面任何一环缺失（无 _origin.pdf / 无 Fig 标题 / pypdfium2 不可用），退回旧逻辑：
           直接复制 mineru 切好的 img_path，按 caption 命名 `Fig.N.jpg` 或 `image_{i}.jpg`。
        """
        try:
            self.log("  - 提取图片...")

            content_list_file = self.find_file_by_glob(
                raw_dir, f"{pdf_name}_content_list.json", "*_content_list.json"
            )
            middle_json_file = self.find_file_by_glob(
                raw_dir, f"{pdf_name}_middle.json", "*_middle.json"
            )
            origin_pdf_file = self.find_file_by_glob(
                raw_dir, f"{pdf_name}_origin.pdf", "*_origin.pdf"
            )

            if not content_list_file:
                self.log("    ⚠️  未找到 content_list.json")
                return False

            with open(content_list_file, 'r', encoding='utf-8') as f:
                content_list = json.load(f)

            fig_caption_pat = re.compile(
                r'\b(Fig\.?|Figure|Scheme|图|示意图|插图)\s*(\d+)', re.IGNORECASE
            )

            # 按阅读顺序分组：逐页遍历 content_list 里的 image/chart 条目，
            # 每个带 "Fig N" caption 的条目 claim 自己 + 上一个 caption 以来所有未标号的 pending 子面板。
            # 尾部未标号的塞给该页最后一个 Fig。
            def _extract_caption(it):
                caps = (
                    it.get('image_caption')
                    or it.get('chart_caption')
                    or it.get('img_caption')
                    or []
                )
                return ' '.join(caps).strip() if caps else ''

            # 先收集每页的 image/chart 条目（按 content_list 顺序）+ 对应 content_list 索引
            page_items = {}  # page_idx -> [(cl_idx, item, caption_text, fig_num_or_None)]
            fig_entries = []  # 仅 Fig-captioned 条目（给回退路径用）
            for cl_idx, item in enumerate(content_list):
                if not isinstance(item, dict) or item.get('type') not in ('image', 'chart'):
                    continue
                page_idx = item.get('page_idx')
                if page_idx is None:
                    continue
                caption = _extract_caption(item)
                m = fig_caption_pat.search(caption) if caption else None
                fig_num = m.group(2) if m else None
                page_items.setdefault(page_idx, []).append((cl_idx, item, caption, fig_num))
                if fig_num:
                    fig_entries.append({
                        'fig_num': fig_num,
                        'page_idx': page_idx,
                        'bbox': item.get('bbox'),
                        'img_path': item.get('img_path'),
                        'caption_text': caption,
                    })

            # 对每页做阅读顺序分组
            # groups_per_page[page_idx] = [(fig_num, [cl_indices])]
            groups_per_page = {}
            for p_idx, items in page_items.items():
                pending = []      # 未标号的 cl_indices
                groups = []       # [(fig_num, [cl_indices])]
                for cl_idx, item, caption, fig_num in items:
                    if fig_num:
                        claimed = pending + [cl_idx]
                        groups.append((fig_num, claimed))
                        pending = []
                    else:
                        pending.append(cl_idx)
                # 尾部未标号的塞给最后一个 Fig；若本页没 Fig 则丢弃
                if pending and groups:
                    last_fig, last_idxs = groups[-1]
                    groups[-1] = (last_fig, last_idxs + pending)
                if groups:
                    groups_per_page[p_idx] = groups

            figure_dir = extract_dir / "Figure"
            figure_dir.mkdir(exist_ok=True)
            image_mapping = {}
            saved_figs = []
            sub_img_to_fig = {}

            # ---- 路径 1：重渲染 ----
            can_render = bool(middle_json_file and origin_pdf_file and fig_entries)
            middle = None
            pdf_info = []
            if can_render:
                try:
                    with open(middle_json_file, 'r', encoding='utf-8') as f:
                        middle = json.load(f)
                    pdf_info = middle.get('pdf_info', [])
                except Exception as e:
                    self.log(f"    ⚠️  middle.json 读取失败，回退到复制模式: {e}")
                    can_render = False

            if can_render:
                # 预计算每页按 index 排好序的 middle.json image/chart block 列表
                # （与 content_list 每页 image/chart 条目顺序对齐，位置 k→k 映射）
                page_blocks_ordered = {}
                for p_idx in range(len(pdf_info)):
                    pb = pdf_info[p_idx].get('para_blocks') or []
                    img_blocks = [
                        b for b in pb
                        if isinstance(b, dict)
                        and b.get('type') in ('image', 'chart')
                        and b.get('bbox')
                    ]
                    img_blocks.sort(key=lambda b: b.get('index', 0))
                    page_blocks_ordered[p_idx] = img_blocks

                # 将每个 cl_idx 映射到该页 content_list image/chart 顺序中的位置
                # 然后再映射到 middle.json 同位置的 block bbox
                for p_idx, groups in sorted(groups_per_page.items()):
                    page_item_list = [t[0] for t in page_items[p_idx]]  # 本页 cl_idx 顺序
                    img_blocks = page_blocks_ordered.get(p_idx, [])
                    if not img_blocks:
                        continue
                    # 按 Fig 号排序输出（Fig.1、Fig.2、…）
                    for fig_num, cl_indices in sorted(groups, key=lambda g: int(g[0])):
                        bboxes = []
                        for ci in cl_indices:
                            try:
                                pos = page_item_list.index(ci)
                            except ValueError:
                                continue
                            if pos < len(img_blocks):
                                bboxes.append(img_blocks[pos]['bbox'])
                        bbox = _union_bbox(bboxes)
                        if not bbox:
                            continue
                        out_path = figure_dir / f"Fig.{fig_num}.jpg"
                        try:
                            img = _render_pdf_region(origin_pdf_file, p_idx, bbox, dpi=200)
                            img.save(out_path, "JPEG", quality=92)
                            saved_figs.append(fig_num)
                            # 该组内所有条目的 img_path 都指向这张 Fig.N.jpg
                            for ci in cl_indices:
                                it = content_list[ci] if ci < len(content_list) else None
                                if isinstance(it, dict):
                                    ip = it.get('img_path')
                                    if ip:
                                        sub_img_to_fig[ip] = fig_num
                            self.log(
                                f"    ✓ Fig.{fig_num} (page {p_idx+1}, "
                                f"{img.size[0]}×{img.size[1]}px, {len(bboxes)} 个面板)"
                            )
                        except Exception as e:
                            self.log(f"    ⚠️  Fig.{fig_num} 重渲染失败: {e}")

                for ip, fig_num in sub_img_to_fig.items():
                    if fig_num in saved_figs:
                        image_mapping[ip] = f"Fig.{fig_num}.jpg"

            # ---- 路径 2：回退到复制模式（缺失原 PDF 或 middle.json）----
            if not saved_figs:
                if can_render:
                    self.log("    ⚠️  重渲染未产出，回退到复制模式")
                for i, entry in enumerate(fig_entries, 1):
                    if not entry.get('img_path'):
                        continue
                    src = raw_dir / entry['img_path']
                    if not src.exists():
                        self.log(f"    ⚠️  图片不存在: {src}")
                        continue
                    fig_num = entry.get('fig_num') or str(i)
                    dst = figure_dir / f"Fig.{fig_num}{src.suffix}"
                    try:
                        shutil.copy2(src, dst)
                        saved_figs.append(fig_num)
                        image_mapping[entry['img_path']] = dst.name
                    except Exception as e:
                        self.log(f"    ⚠️  复制 Fig.{fig_num} 失败: {e}")

            mapping_file = figure_dir / "image_mapping.json"
            with open(mapping_file, 'w', encoding='utf-8') as f:
                json.dump(image_mapping, f, ensure_ascii=False, indent=2)

            if saved_figs:
                uniq = sorted(set(saved_figs), key=lambda x: int(x) if x.isdigit() else 9999)
                self.log(f"    ✓ 共生成 {len(uniq)} 张主图: {', '.join('Fig.' + n for n in uniq)}")
                return True
            self.log("    ⚠️  未生成任何图片")
            return False

        except Exception as e:
            self.log(f"    ❌ 图片提取失败: {e}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def extract_tables(self, raw_dir: Path, extract_dir: Path, pdf_name: str) -> bool:
        """提取表格（从 content_list.json，保存为 Excel 和 JPG）"""
        try:
            self.log("  - 提取表格...")

            # 读取 content_list.json
            content_list_file = self.find_file_by_glob(raw_dir, f"{pdf_name}_content_list.json", "*_content_list.json")
            if content_list_file is None:
                self.log("    ⚠️  未找到 content_list.json 文件")
                return False

            with open(content_list_file, 'r', encoding='utf-8') as f:
                content_list = json.load(f)

            # 提取所有表格
            tables = []
            for item in content_list:
                if item.get('type') == 'table':
                    tables.append(item)

            if not tables:
                self.log("    ⚠️  未找到表格")
                return False

            self.log(f"    找到 {len(tables)} 个表格")

            # 创建 Tables 文件夹（与 Figure、Formula 同级）
            tables_dir = extract_dir / "Tables"
            tables_dir.mkdir(exist_ok=True)

            # 保存表格
            for i, table_item in enumerate(tables, 1):
                # 获取表格标题和编号
                table_caption = ""
                table_number = str(i)

                if 'table_caption' in table_item and table_item['table_caption']:
                    caption_text = latex_to_unicode(' '.join(table_item['table_caption']))
                    table_caption = caption_text

                    # 尝试从标题中提取表格编号
                    caption_match = re.search(r'Table\s*(\d+)', caption_text, re.IGNORECASE)
                    if caption_match:
                        table_number = caption_match.group(1)

                # 1. 保存表格图片（始终保存）
                if 'img_path' in table_item:
                    img_path = raw_dir / table_item['img_path']
                    if img_path.exists():
                        img_dest = tables_dir / f"Table_{table_number}{img_path.suffix}"
                        shutil.copy2(img_path, img_dest)
                        self.log(f"    ✓ 保存表格图片 {table_number}: {img_dest.name}")

                # 2. 保存表格为 Excel（Sheet 名称为文件名，表头显示完整标题）
                try:
                    import pandas as pd
                    from bs4 import BeautifulSoup

                    if 'table_body' in table_item:
                        # 解析 HTML 表格
                        soup = BeautifulSoup(table_item['table_body'], 'html.parser')
                        table = soup.find('table')

                        if table:
                            # 提取表格数据；mineru 3.x 在单元格里保留 LaTeX 源码
                            # （如 "$\sigma ( \mathsf{S}\mathsf{cm} ^ { -1 } )$"），
                            # 这里统一转成 Unicode 让 Excel 里直接可读。
                            # 同时修复 3.x 偶发的 "行首空格 + 次列塞两行 \begin{array}" 错位
                            # （例："Simulated" 行标签被吞进数据列的 array 第 1 行）。
                            array_split_pat = re.compile(
                                r'^\$\s*\\begin\{array\}\s*\{[^}]*\}\s*(.*?)\s*\\end\{array\}\s*\$\s*$',
                                re.DOTALL,
                            )
                            # row 2 若形如 "<unit>\right) <value>" （label 后半 + 数值拼一起），
                            # 按 \right) 再切一刀：前半合入 A 列 label，后半留 B 列 value。
                            right_split_pat = re.compile(
                                r'^(.*?\\right\s*\)\s*)(.+)$', re.DOTALL
                            )
                            rows = []
                            for tr in table.find_all('tr'):
                                raw = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                                if len(raw) >= 2 and raw[0] == '':
                                    m = array_split_pat.match(raw[1])
                                    if m:
                                        parts = re.split(r'\s*\\\\\s*', m.group(1))
                                        if len(parts) == 2:
                                            r1 = re.sub(r'^\s*\{+\s*|\s*\}+\s*$', '', parts[0].strip())
                                            r2 = re.sub(r'^\s*\{+\s*|\s*\}+\s*$', '', parts[1].strip())
                                            rs = right_split_pat.match(r2)
                                            if rs:
                                                # label = r1 + r2 前缀（到 \right)）；两段分别包 $...$ 保留中间空格
                                                rs_prefix = rs.group(1).strip()
                                                value_part = rs.group(2).strip()
                                                raw[0] = f"${r1}$ ${rs_prefix}$" if r1 and rs_prefix else (
                                                    f"${r1}$" if r1 else (f"${rs_prefix}$" if rs_prefix else '')
                                                )
                                                raw[1] = f"${value_part}$" if value_part else ''
                                            else:
                                                raw[0] = f"${r1}$" if r1 else ''
                                                raw[1] = f"${r2}$" if r2 else ''
                                rows.append([latex_to_unicode(c) for c in raw])

                            # 创建 DataFrame
                            if rows:
                                # 方法1：尝试使用第一行作为表头
                                try:
                                    df = pd.DataFrame(rows[1:], columns=rows[0] if len(rows) > 1 else None)
                                except ValueError as e:
                                    # 列数不匹配（通常是因为 rowspan/colspan）
                                    # 方法2：找到列数最多的行作为参考
                                    max_cols = max(len(row) for row in rows)
                                    
                                    # 尝试使用列数最多的行作为参考，不使用表头
                                    self.log(f"    ⓘ 表格 {table_number} 结构复杂（rowspan/colspan），使用备用方案")
                                    
                                    # 统一列数：不足的行用空字符串填充
                                    normalized_rows = []
                                    for row in rows:
                                        if len(row) < max_cols:
                                            row = row + [''] * (max_cols - len(row))
                                        normalized_rows.append(row[:max_cols])  # 截取到最大列数
                                    
                                    # 不使用表头，使用默认列名（Column 1, Column 2, ...）
                                    df = pd.DataFrame(normalized_rows)

                                # Sheet 名称使用文件名（Table_1）
                                sheet_name = f"Table_{table_number}"

                                # 保存为 Excel
                                excel_file = tables_dir / f"Table_{table_number}.xlsx"

                                # 使用 ExcelWriter 来添加表格标题
                                with pd.ExcelWriter(excel_file, engine='openpyxl') as writer:
                                    df.to_excel(writer, index=False, sheet_name=sheet_name)

                                    # 获取工作表
                                    worksheet = writer.sheets[sheet_name]

                                    # 在第一行插入表格标题（如果有）
                                    if table_caption:
                                        worksheet.insert_rows(1)
                                        worksheet['A1'] = table_caption
                                        # 合并第一行的单元格
                                        from openpyxl.styles import Font, Alignment
                                        worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(df.columns))
                                        worksheet['A1'].font = Font(bold=True, size=12)
                                        worksheet['A1'].alignment = Alignment(horizontal='center', vertical='center')

                                self.log(f"    ✓ 保存表格 Excel {table_number}: {excel_file.name}")

                except ImportError:
                    self.log("    ⚠️  需要安装 pandas 和 openpyxl 来生成 Excel 文件")
                    self.log("    ⚠️  运行: pip install pandas openpyxl beautifulsoup4")
                except Exception as e:
                    self.log(f"    ⚠️  表格 {table_number} Excel 生成失败: {str(e)}")

            self.log(f"    ✓ 提取 {len(tables)} 个表格")
            return True

        except Exception as e:
            self.log(f"    ❌ 表格提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())
            return False

    def extract_sections_with_llm(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """提取论文章节（优先使用正则表达式，失败时使用 LLM）"""
        try:
            self.log("  - 提取论文章节...")

            # 检查 LLM 模块是否可用
            if not LLM_AVAILABLE:
                self.log("    ❌ LLM 模块不可用，请检查 llm_helper.py 是否存在")
                return

            # 读取 Markdown 文件
            md_file = self.find_file_by_glob(extract_dir, f"{pdf_name}.md", "*.md")
            if md_file is None:
                # 尝试从 raw 目录读取
                md_file = self.find_file_by_glob(raw_dir, f"{pdf_name}.md", "*.md")

            if md_file is None:
                self.log("    ❌ 未找到 Markdown 文件，请先提取文字")
                return

            self.log("    读取 Markdown 文件...")
            with open(md_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # 检查文件大小（避免超过 API 限制）
            content_length = len(markdown_content)
            self.log(f"    文档长度: {content_length} 字符")

            if content_length > 100000:
                self.log("    ⚠️  文档过长，可能超过 API 限制，尝试截取前 100000 字符")
                markdown_content = markdown_content[:100000]

            # 提示词模板将在需要时加载（根据是否有缺失章节选择不同模板）
            prompt_template = None

            # 初始化 LLM
            model_name = self._run_option("llm_model", "")
            provider = self._run_option("llm_provider", "deepseek")
            provider_label = "DeepSeek" if provider == "deepseek" else "自定义接口"
            self.log(f"    使用接口: {provider_label}")
            self.log(f"    使用模型: {model_name}")

            try:
                llm = LLMHelper(
                    model_name=model_name,
                    provider=provider,
                    env_path=self.base_path / ".env",
                    require_api=False,
                    debug_dir=self.output_path / "debug",
                )
            except ValueError as e:
                self.log(f"    ❌ LLM 初始化失败: {str(e)}")
                self.log("    💡 请在“设置”中测试接口、选择模型并保存")
                return

            # 优先使用正则表达式提取章节（同时获取未识别的标题）
            self.log("    🔄 使用正则表达式提取章节...")
            sections, unrecognized_headers = llm.extract_sections_fallback(markdown_content, return_unrecognized=True)

            # 质量检查：判断是否需要使用 LLM
            need_llm = False
            llm_reason = []

            if not sections:
                # 情况1: 正则表达式完全失败
                need_llm = True
                llm_reason.append("正则表达式未能提取任何章节")
            else:
                # 情况2: 检查关键章节是否缺失
                critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                missing_critical = [s for s in critical_sections if s not in sections]
                
                if missing_critical:
                    need_llm = True
                    llm_reason.append(f"缺少关键章节: {', '.join(missing_critical)}")
                
                # 情况3: 检查章节内容是否过短（可能截断）
                short_sections = []
                for name, content in sections.items():
                    if len(content.strip()) < 100:  # 少于100字符
                        short_sections.append(name)
                
                if short_sections:
                    need_llm = True
                    llm_reason.append(f"章节内容过短: {', '.join(short_sections)}")
                
                # 情况4: 章节数量异常
                if len(sections) < 2:
                    need_llm = True
                    llm_reason.append(f"章节数量过少 ({len(sections)}个)")
                elif len(sections) > 8:
                    # 章节过多可能是误识别，但不强制使用LLM
                    self.log(f"    ⓘ 提取到较多章节 ({len(sections)}个)，可能包含子章节")

            # 优先尝试使用 LLM 对未识别的标题进行分类（轻量级方案）
            if unrecognized_headers and sections:
                self.log(f"    🔍 发现 {len(unrecognized_headers)} 个未识别的章节标题，尝试使用 LLM 分类...")
                try:
                    classification = llm.classify_section_titles(unrecognized_headers)

                    if classification:
                        # 根据分类结果，从 markdown 中提取对应章节内容
                        lines = markdown_content.split('\n')

                        # 按章节类型分组（处理多个标题映射到同一章节类型的情况）
                        sections_by_type = {}
                        for header_text, section_type in classification.items():
                            if section_type not in sections_by_type:
                                sections_by_type[section_type] = []
                            sections_by_type[section_type].append(header_text)

                        # 提取每个章节类型的内容
                        for section_type, headers in sections_by_type.items():
                            # 检查该章节是否已经被正则提取过
                            already_exists = section_type in sections
                            if already_exists:
                                self.log(f"       ⓘ {section_type} 已存在（正则提取），将合并 LLM 分类的内容")

                            # 提取所有匹配该类型的标题的内容，并合并
                            combined_content = []

                            # 找到所有标题在文档中的位置
                            header_positions = []
                            for header_text in headers:
                                for line_idx, line in enumerate(lines):
                                    if line.strip() == header_text:
                                        header_positions.append((line_idx, header_text))
                                        break

                            # 按位置排序
                            header_positions.sort(key=lambda x: x[0])

                            # 提取内容：从第一个标题到最后一个标题之后的下一个同级标题
                            if header_positions:
                                start_idx = header_positions[0][0]
                                last_idx = header_positions[-1][0]

                                # 找到章节结束位置（下一个同级或更高级的标题）
                                # 获取起始标题的级别（例如 "# 2." 是一级标题）
                                start_line = lines[start_idx].strip()
                                import re
                                start_match = re.match(r'^(#+)\s+(\d+)\.', start_line)
                                if start_match:
                                    start_level_hashes = len(start_match.group(1))
                                    start_number = int(start_match.group(2))

                                    # 从最后一个标题之后开始查找
                                    section_end = len(lines)
                                    for i in range(last_idx + 1, len(lines)):
                                        line_stripped = lines[i].strip()
                                        # 检查是否是同级或更高级的标题
                                        # 匹配 "# 数字." 或 "# 数字.数字." 等格式
                                        match = re.match(r'^(#+)\s+(\d+)(?:\.(\d+))*\.', line_stripped)
                                        if match:
                                            level_hashes = len(match.group(1))
                                            number = int(match.group(2))
                                            sub_number = match.group(3)  # 子编号（如 3.1 中的 1）

                                            # 判断是否是同级或更高级标题：
                                            # 1. 如果 hash 数量相同且没有子编号，且主编号更大 → 同级标题
                                            # 2. 如果 hash 数量更少 → 更高级标题
                                            if level_hashes == start_level_hashes:
                                                # 同级标题：没有子编号，且主编号更大
                                                if sub_number is None and number > start_number:
                                                    section_end = i
                                                    break
                                            elif level_hashes < start_level_hashes:
                                                # 更高级标题
                                                section_end = i
                                                break
                                        # 也检查是否是排除章节（Acknowledgements, References 等）
                                        elif line_stripped.startswith('# ') and not re.match(r'^#+\s+\d+\.', line_stripped):
                                            section_end = i
                                            break
                                else:
                                    # 如果无法解析标题级别，使用简单逻辑：找下一个一级标题
                                    section_end = len(lines)
                                    for i in range(last_idx + 1, len(lines)):
                                        if lines[i].strip().startswith('# ') and not lines[i].strip().startswith('## '):
                                            section_end = i
                                            break

                                # 提取完整内容
                                section_content = '\n'.join(lines[start_idx:section_end])
                                combined_content.append(section_content)
                                self.log(f"       ✓ 提取章节内容: {headers[0][:40]}... 到 {headers[-1][:40]}... → {section_type}")

                            # 合并所有内容
                            if combined_content:
                                new_content = '\n\n'.join(combined_content)
                                if already_exists:
                                    # 如果章节已存在，将新内容添加到前面（因为 LLM 分类的通常是前面的章节）
                                    sections[section_type] = new_content + '\n\n' + sections[section_type]
                                    self.log(f"       ✓ 合并到现有章节: {section_type} (添加 {len(combined_content)} 个片段)")
                                else:
                                    sections[section_type] = new_content
                                    self.log(f"       ✓ 通过分类补充章节: {section_type} (合并 {len(combined_content)} 个片段)")

                        # 重新检查是否还有缺失的关键章节
                        critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                        missing_critical = [s for s in critical_sections if s not in sections]
                        if missing_critical:
                            self.log(f"    ⓘ 分类后仍缺少: {', '.join(missing_critical)}")
                        else:
                            self.log(f"    ✅ 通过标题分类成功补全所有关键章节！")
                            need_llm = False  # 不需要再用 LLM 提取全文了

                except Exception as e:
                    self.log(f"    ⚠️  标题分类失败: {str(e)}")

            # 如果标题分类后仍有问题，使用 LLM 补充或重新提取
            if need_llm:
                # Phase 10: 若论文明显不符合 IEEE 四段式模板（正则只抓到 < 2 个章节且缺 ≥ 3 个关键章节），
                # 跳过昂贵的 LLM 补全。Nature / Science 风格论文常见。继续调 LLM 基本只会生成伪造内容。
                critical = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                missing_critical = [s for s in critical if s not in sections]
                if len(sections) < 2 and len(missing_critical) >= 3:
                    self.log(
                        f"    ⓘ 论文缺少标准章节结构（正则仅识别 {len(sections)} 个，缺 {len(missing_critical)} 个），跳过 LLM 补全"
                    )
                    self.log("    📝 原文可能是 Nature/Science 风格，没有显式 Abstract/Introduction/Results 标题")
                    self.log(f"    ↳ 如需全文请看 {pdf_name}.md；LLM 生成的内容易脱离原文，故不强行补")
                    need_llm = False

            if need_llm:
                self.log("    ⚠️  质量检查发现问题:")
                for reason in llm_reason:
                    self.log(f"       - {reason}")
                self.log("    🤖 尝试使用 LLM 改进提取结果...")
                self.log("    ⏳ 这可能需要 10-30 秒，请耐心等待...")

                try:
                    # 确定需要提取的章节
                    if not sections:
                        # 如果正则完全失败，提取所有章节
                        missing_sections = None
                        prompt_file = Path(__file__).parent / "prompts" / "section_extraction_prompt.txt"
                        self.log("    📋 正则提取失败，使用 LLM 提取所有章节")
                    else:
                        # 只提取缺失的章节
                        critical_sections = ['Abstract', 'Introduction', 'Methods', 'Results & Discussion', 'Conclusion']
                        missing_sections = [s for s in critical_sections if s not in sections]

                        # 同时包含过短的章节
                        for name, content in sections.items():
                            if len(content.strip()) < 100 and name not in missing_sections:
                                missing_sections.append(name)

                        if missing_sections:
                            prompt_file = Path(__file__).parent / "prompts" / "section_extraction_missing_prompt.txt"
                            self.log(f"    📋 只提取缺失章节: {', '.join(missing_sections)}")
                        else:
                            # 没有缺失章节，不需要调用 LLM
                            self.log("    ⓘ 没有缺失章节，跳过 LLM 调用")
                            need_llm = False

                    if need_llm:
                        # 加载提示词模板
                        if not prompt_file.exists():
                            self.log(f"    ❌ 提示词文件不存在: {prompt_file}")
                        else:
                            self.log("    加载提示词模板...")
                            prompt_template = load_prompt_template(prompt_file)

                            llm_sections = llm.extract_sections(markdown_content, prompt_template, missing_sections)

                            if llm_sections:
                                # 合并策略：优先使用LLM结果，但保留正则表达式的优质结果
                                if not sections:
                                    # 如果正则完全失败，直接使用LLM结果
                                    sections = llm_sections
                                    self.log(f"    ✓ LLM 成功提取 {len(llm_sections)} 个章节")
                                else:
                                    # 智能合并：补充缺失的章节，替换过短的章节
                                    merged_count = 0
                                    for name, llm_content in llm_sections.items():
                                        if name not in sections:
                                            # 补充缺失的章节
                                            sections[name] = llm_content
                                            merged_count += 1
                                            self.log(f"       ✓ 补充章节: {name}")
                                        elif len(sections[name].strip()) < 100 and len(llm_content.strip()) > 100:
                                            # 用LLM的更完整内容替换过短的章节
                                            sections[name] = llm_content
                                            merged_count += 1
                                            self.log(f"       ✓ 改进章节: {name}")

                                    if merged_count > 0:
                                        self.log(f"    ✓ 成功合并 {merged_count} 个章节")
                                    else:
                                        self.log(f"    ⓘ LLM 结果未提供改进")
                            else:
                                self.log("    ⚠️  LLM 提取失败")

                except Exception as e:
                    self.log(f"    ⚠️  LLM 调用出错: {str(e)}")
                    self.log("    ℹ️  将继续使用正则表达式的结果")

            if not sections:
                self.log("    ❌ 所有方法都失败了，无法提取章节")
                self.log("    💡 建议：检查文档格式，或手动提取章节")
                return

            self.log(f"    ✓ 成功识别到 {len(sections)} 个章节")

            # 保存章节
            sections_dir = extract_dir / "Sections"
            saved_files = save_sections(sections, sections_dir)

            self.log(f"    ✓ 保存了 {len(saved_files)} 个章节文件:")
            for file_path in saved_files:
                file_name = Path(file_path).name
                self.log(f"      - {file_name}")

        except Exception as e:
            self.log(f"    ❌ 章节提取失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    def create_word_folder(self, raw_dir: Path, extract_dir: Path, pdf_name: str):
        """创建 Word 文档和 Markdown 图表汇总（按原文顺序排列图片和表格）"""
        try:
            self.log("  - 创建 Word 文档和 Markdown 图表汇总...")

            # 检查是否安装了 python-docx
            try:
                from docx import Document
                from docx.shared import Cm, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                self.log("    ⚠️  需要安装 python-docx 来生成 Word 文档")
                self.log("    ⚠️  运行: pip install python-docx")
                return

            # 读取 Markdown 文件以获取图表标题
            md_file = self.find_file_by_glob(raw_dir, f"{pdf_name}.md", "*.md")
            if md_file is None:
                self.log("    ⚠️  未找到 Markdown 文件")
                return

            with open(md_file, 'r', encoding='utf-8') as f:
                md_lines = f.readlines()

            # 创建 Word 文档
            doc = Document()

            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.left_margin = Cm(2.54)
                section.right_margin = Cm(2.54)

            # 创建 Word 文件夹
            word_dir = extract_dir / "Word"
            word_dir.mkdir(exist_ok=True)

            # 获取 Tables 文件夹路径
            tables_dir = extract_dir / "Tables"
            figure_dir = extract_dir / "Figure"

            # 读取图片映射文件
            image_mapping = {}
            mapping_file = figure_dir / "image_mapping.json"
            if mapping_file.exists():
                try:
                    with open(mapping_file, 'r', encoding='utf-8') as f:
                        image_mapping = json.load(f)
                except Exception as e:
                    self.log(f"    ⚠️  读取图片映射文件失败: {str(e)}")

            # 准备 Markdown 图表汇总内容
            md_summary_lines = []
            md_summary_lines.append(f"# {pdf_name} - 图表汇总\n\n")
            md_summary_lines.append("本文档包含从 PDF 中提取的所有图片和表格，按原文顺序排列。\n\n")
            md_summary_lines.append("---\n\n")

            # 遍历 Markdown 文件，按顺序处理图片和表格
            i = 0
            item_count = 0
            # mineru 3.x 会把一张 Figure 切成多个子面板，主 .md 里同一个 Fig N 会被重复引用；
            # 记录已插入的 Fig 编号避免重复 add_picture。
            inserted_fig_nums = set()

            while i < len(md_lines):
                line = md_lines[i].strip()

                # 检查是否是表格标题（表格标题在上方）
                if re.match(r'^(Table|表)\s*\d+', line) or (line.startswith('#') and re.search(r'(Table|表)\s*\d+', line)):
                    # 去掉可能存在的#号和空格
                    title = re.sub(r'^#+\s*', '', line)

                    # 添加表格标题（表格标题在表上方）
                    p = doc.add_paragraph(title)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_format = p.paragraph_format
                    p_format.space_before = Pt(6)
                    p_format.space_after = Pt(6)

                    # 从标题中提取表格编号
                    table_match = re.search(r'(Table|表)\s*(\d+)', title, re.IGNORECASE)
                    if table_match:
                        table_number = table_match.group(2)

                        # 直接从 Tables 文件夹中查找表格图片
                        table_img = tables_dir / f"Table_{table_number}.jpg"
                        if not table_img.exists():
                            table_img = tables_dir / f"Table_{table_number}.png"

                        if table_img.exists():
                            # 添加表格图片到 Word
                            doc.add_picture(str(table_img), width=Cm(14))
                            doc.add_paragraph()  # 添加空行

                            # 添加到 Markdown 汇总
                            md_summary_lines.append(f"## {title}\n\n")
                            # 使用相对路径引用图片
                            relative_path = f"../Tables/{table_img.name}"
                            md_summary_lines.append(f"![{title}]({relative_path})\n\n")
                            md_summary_lines.append("---\n\n")

                            item_count += 1
                            self.log(f"      添加表格 {table_number}")
                        else:
                            self.log(f"      ⚠️  未找到表格图片: {table_img.name}")

                    i += 1
                    continue

                # 检查是否是图片（每个 Fig 号只插入一次；优先用 image_mapping.json 决定 Fig 号）
                if line.startswith('!['):
                    img_match = re.search(r'\((.*?)\)', line)
                    if img_match:
                        img_path = img_match.group(1).strip()
                        img_full_path = raw_dir / img_path

                        # 是否是表格图片（表格由上面的分支单独处理）
                        is_table_image = False
                        for j in range(max(0, i - 5), i):
                            if re.match(r'^(Table|表)\s*\d+', md_lines[j].strip()):
                                is_table_image = True
                                break

                        if not is_table_image:
                            # 主路径：通过 image_mapping.json 找到对应 Fig.N（extract_figures 已做好阅读顺序分组）
                            fig_num = None
                            caption = ""
                            mapped = image_mapping.get(img_path) or image_mapping.get(img_path.rstrip())
                            if mapped:
                                mm = re.match(r'Fig\.(\d+)\.jpg', mapped, re.IGNORECASE)
                                if mm:
                                    fig_num = mm.group(1)

                            # 回退路径（兼容 pre-Phase-12 输出）：下一行是 "Fig N" 标题
                            if not fig_num and i + 1 < len(md_lines):
                                nxt = md_lines[i + 1].strip()
                                cap_m = re.match(
                                    r'^(Fig\.?|Figure|图|Scheme|图表|示意图|插图)\s*(\d+)',
                                    nxt, re.IGNORECASE
                                )
                                if cap_m:
                                    caption = nxt
                                    fig_num = cap_m.group(2)

                            # 给通过 mapping 发现的 Fig 找标题。
                            # 真图题的特征：**行首**就是 "Fig N." 并后接一段有实际长度的描述；
                            # 正文里 "...as shown in Fig. 4." 这种中间引用句要排除，否则会把正文
                            # 句子当作图题，或者被截断到只剩 "Fig. 4." 的空壳。
                            if fig_num and not caption:
                                # 行首形式 + 必须有 >=3 字符的描述内容（去掉尾点后）
                                real_title_pat = re.compile(
                                    r'^\s*(?:Fig\.?|Figure|图|Scheme|图表|示意图|插图)\s*'
                                    + re.escape(fig_num)
                                    + r'\s*[.:：、]?\s*(.+)$',
                                    re.IGNORECASE,
                                )

                                def _is_real_title(line: str) -> bool:
                                    m = real_title_pat.match(line.strip())
                                    if not m:
                                        return False
                                    rest = m.group(1).strip().rstrip('.。').strip()
                                    return len(rest) >= 3

                                # Step 1: 图片紧随其后 12 行内找真图题（mineru 切子面板时，
                                # 真图题常隔着几行 "(a) xxx" / "(b) xxx" 才出现）
                                for jj in range(i + 1, min(len(md_lines), i + 13)):
                                    if _is_real_title(md_lines[jj]):
                                        caption = md_lines[jj].strip()
                                        break

                                # Step 2: 退而求其次，全文扫首条真图题行
                                if not caption:
                                    for jj in range(len(md_lines)):
                                        if _is_real_title(md_lines[jj]):
                                            caption = md_lines[jj].strip()
                                            break

                                # Step 3: 兜底（极少数图确实没有图题）
                                if not caption:
                                    caption = f"Fig. {fig_num}."

                            if fig_num and fig_num not in inserted_fig_nums:
                                # 优先使用 extract_figures 重渲染出的完整 Fig.N.jpg；
                                # 若缺失，再退回到 md 里引用的原始（子面板）img
                                rebuilt = figure_dir / f"Fig.{fig_num}.jpg"
                                if rebuilt.exists():
                                    insert_src = rebuilt
                                    relative_path = f"../Figure/{rebuilt.name}"
                                elif img_full_path.exists():
                                    insert_src = img_full_path
                                    if img_path in image_mapping:
                                        relative_path = f"../Figure/{image_mapping[img_path]}"
                                    else:
                                        relative_path = f"../{img_path}"
                                else:
                                    insert_src = None
                                    relative_path = None

                                if insert_src is not None:
                                    inserted_fig_nums.add(fig_num)

                                    doc.add_picture(str(insert_src), width=Cm(14))

                                    # 将 $...$ LaTeX 段转成可读 Unicode（与表格标题一致），
                                    # 避免 Word / md 里出现裸 \mathbf{c}、\sigma 等代码。
                                    caption_display = latex_to_unicode(caption)

                                    p = doc.add_paragraph(caption_display)
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    p_format = p.paragraph_format
                                    p_format.space_before = Pt(6)
                                    p_format.space_after = Pt(12)

                                    md_summary_lines.append(f"## {caption_display}\n\n")
                                    md_summary_lines.append(f"![{caption_display}]({relative_path})\n\n")
                                    md_summary_lines.append("---\n\n")

                                    item_count += 1
                                    i += 2  # 跳过 caption 行
                                    continue

                            # 其它情形（子面板 / 已去重的 Fig / 找不到源文件）：跳过图片不入 docx
                i += 1

            # 保存 Word 文档；若完整路径预计超过 Windows MAX_PATH (260)，
            # 退化到短文件名 "图表.docx" 以保证 Word 能打开（目录名仍保留完整 pdf_name 供导航）。
            # 留 10 字符余量给 ".docx"/"_图表汇总.md" 等后缀。
            default_name = f"{pdf_name}_图表.docx"
            default_path = word_dir / default_name
            if len(str(default_path.resolve())) > 245:
                self.log(f"    ⚠️  完整路径过长 ({len(str(default_path.resolve()))} > 245)，改用短文件名 '图表.docx'")
                output_doc = word_dir / "图表.docx"
            else:
                output_doc = default_path
            doc.save(str(output_doc))
            # python-docx 偶发 bug：保存后 zip 里缺失 numbering.xml 等被 rels/Content_Types 声明的部件，
            # 导致 Word 打开时报"损坏"。在这里自愈一下：把声明了但缺失的 XML 部件补个最小占位。
            _repair_docx_missing_parts(output_doc)

            # 保存 Markdown 图表汇总；同样对长路径做兜底
            default_md_name = f"{pdf_name}_图表汇总.md"
            default_md_path = word_dir / default_md_name
            if len(str(default_md_path.resolve())) > 245:
                md_summary_file = word_dir / "图表汇总.md"
            else:
                md_summary_file = default_md_path
            with open(md_summary_file, 'w', encoding='utf-8') as f:
                f.writelines(md_summary_lines)

            self.log(f"    ✓ Word 文档创建完成: {output_doc.name}")
            self.log(f"    ✓ Markdown 图表汇总创建完成: {md_summary_file.name}")
            self.log(f"    ✓ 包含 {item_count} 个图表")

        except Exception as e:
            self.log(f"    ❌ Word 文档创建失败: {str(e)}")
            import traceback
            self.log(traceback.format_exc())

    # =============================================
    # 环境检测
    # =============================================

    def check_environment(self):
        """首次运行环境检测，检查核心依赖是否已安装"""
        issues = []

        # 只读取发行版元数据，不在 GUI 启动阶段加载 CUDA DLL。
        try:
            from importlib.metadata import version as package_version

            mineru_version = package_version("mineru")
            mineru_parts = tuple(
                int(part) for part in re.findall(r"\d+", mineru_version)[:3]
            )
            if not ((3, 1) <= mineru_parts < (4, 0)):
                issues.append(
                    f"MinerU {mineru_version} 不受支持（需要 >=3.1.0,<4.0；请执行重装）"
                )
        except Exception:
            issues.append("MinerU (mineru[core]) 未安装")

        try:
            package_version("torch")
        except Exception:
            issues.append("PyTorch 未安装")

        # 检查核心依赖
        dep_checks = [
            ("pandas", "pandas"),
            ("openpyxl", "openpyxl"),
            ("bs4", "beautifulsoup4"),
            ("docx", "python-docx"),
            ("lxml", "lxml"),
        ]
        for module_name, package_name in dep_checks:
            try:
                __import__(module_name)
            except ImportError:
                issues.append(f"{package_name} 未安装")

        # MinerU 3.0+ 模型会在首次运行时自动下载，不再需要检查 mineru.json

        if issues:
            self._show_environment_dialog(issues)
        else:
            self.log("✓ 环境检测通过，所有依赖已就绪")

    def _show_environment_dialog(self, issues: list):
        """显示环境问题引导对话框"""
        dialog = tk.Toplevel(self.root)
        dialog.title("环境检测")
        dialog.geometry("520x400")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 520) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 400) // 2
        dialog.geometry(f"+{x}+{y}")

        main_frame = tk.Frame(dialog, bg='white', padx=24, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(
            main_frame,
            text="检测到以下环境问题",
            font=('Microsoft YaHei UI', 13, 'bold'),
            bg='white', fg='#e74c3c'
        ).pack(anchor=tk.W, pady=(0, 12))

        tk.Label(
            main_frame,
            text='部分功能可能无法正常使用，建议运行"一键安装.bat"修复：',
            font=('Microsoft YaHei UI', 9),
            bg='white', fg='#666'
        ).pack(anchor=tk.W, pady=(0, 8))

        # 问题列表
        issues_frame = tk.Frame(main_frame, bg='#fff3e0', relief='solid', borderwidth=1)
        issues_frame.pack(fill=tk.X, pady=(0, 16))

        for issue in issues:
            tk.Label(
                issues_frame,
                text=f"  \u2022  {issue}",
                font=('Microsoft YaHei UI', 9),
                bg='#fff3e0', fg='#e65100',
                anchor=tk.W
            ).pack(fill=tk.X, padx=12, pady=3)

        # 提示信息
        tk.Label(
            main_frame,
            text="解决方法：",
            font=('Microsoft YaHei UI', 10, 'bold'),
            bg='white', fg='#333'
        ).pack(anchor=tk.W, pady=(0, 6))

        tk.Label(
            main_frame,
            text='1. 双击项目根目录的"一键安装.bat"自动修复\n'
                 '2. 或参考 docs/快速安装指南.md 手动安装',
            font=('Microsoft YaHei UI', 9),
            bg='white', fg='#555',
            justify=tk.LEFT
        ).pack(anchor=tk.W, pady=(0, 16))

        # 按钮区域
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(fill=tk.X)

        def open_install_guide():
            guide_path = self.base_path / "docs" / "快速安装指南.md"
            if guide_path.exists():
                os.startfile(str(guide_path))
            else:
                messagebox.showinfo("提示", "安装指南文件不存在")

        tk.Button(
            btn_frame,
            text="打开安装指南",
            font=('Microsoft YaHei UI', 9),
            bg='#f5f5f5', fg='#333',
            relief='solid', borderwidth=1,
            cursor='hand2',
            command=open_install_guide
        ).pack(side=tk.LEFT, padx=(0, 8))

        tk.Button(
            btn_frame,
            text="我知道了，继续使用",
            font=('Microsoft YaHei UI', 9),
            bg=self.accent_color, fg='white',
            activebackground='#1a5cc7',
            activeforeground='white',
            relief='flat',
            cursor='hand2',
            command=dialog.destroy
        ).pack(side=tk.RIGHT)

        # 记录日志
        for issue in issues:
            self.log(f"  ⚠ {issue}")

    # =============================================
    # 设置面板
    # =============================================

    def _open_settings_legacy(self):
        """打开设置对话框（API Key 配置）"""
        dialog = tk.Toplevel(self.root)
        dialog.title("设置")
        dialog.geometry("500x340")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        # 居中显示
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 500) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 340) // 2
        dialog.geometry(f"+{x}+{y}")

        main_frame = tk.Frame(dialog, bg='white', padx=24, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        tk.Label(
            main_frame,
            text="API 设置",
            font=('Microsoft YaHei UI', 13, 'bold'),
            bg='white', fg='#333'
        ).pack(anchor=tk.W, pady=(0, 16))

        # DeepSeek API Key
        tk.Label(
            main_frame,
            text="DeepSeek API Key",
            font=('Microsoft YaHei UI', 10),
            bg='white', fg='#333'
        ).pack(anchor=tk.W, pady=(0, 4))

        tk.Label(
            main_frame,
            text="用于智能章节提取，不配置也可使用（正则模式，90%成功率）\n"
                 "配置后成功率提升至 98%。获取地址：https://platform.deepseek.com/",
            font=('Microsoft YaHei UI', 8),
            bg='white', fg='#999',
            justify=tk.LEFT
        ).pack(anchor=tk.W, pady=(0, 8))

        # 输入框区域
        key_frame = tk.Frame(main_frame, bg='white')
        key_frame.pack(fill=tk.X, pady=(0, 12))

        # 读取当前 API Key
        current_key = self._read_api_key()

        key_var = tk.StringVar(value=current_key)
        show_key_var = tk.BooleanVar(value=False)

        key_entry = tk.Entry(
            key_frame,
            textvariable=key_var,
            font=('Consolas', 10),
            show='*',
            relief='solid',
            borderwidth=1
        )
        key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 8))

        def toggle_show():
            key_entry.config(show='' if show_key_var.get() else '*')

        tk.Checkbutton(
            key_frame,
            text="显示",
            variable=show_key_var,
            command=toggle_show,
            bg='white',
            font=('Microsoft YaHei UI', 8)
        ).pack(side=tk.RIGHT)

        # 状态标签
        status_label = tk.Label(
            main_frame,
            text="",
            font=('Microsoft YaHei UI', 9),
            bg='white'
        )
        status_label.pack(anchor=tk.W, pady=(0, 16))

        # 按钮区域
        btn_frame = tk.Frame(main_frame, bg='white')
        btn_frame.pack(fill=tk.X, pady=(8, 0))

        def test_connection():
            api_key = key_var.get().strip()
            if not api_key or api_key == "sk-your_deepseek_api_key_here":
                status_label.config(text="请先输入有效的 API Key", fg='#e74c3c')
                return
            status_label.config(text="正在测试连接...", fg='#f57c00')

            def _do_test():
                try:
                    import requests as req
                    headers = {
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    }
                    data = {
                        "model": "deepseek-chat",
                        "messages": [{"role": "user", "content": "hi"}],
                        "max_tokens": 5
                    }
                    resp = req.post(
                        "https://api.deepseek.com/v1/chat/completions",
                        headers=headers, json=data, timeout=15
                    )
                    if resp.status_code == 200:
                        dialog.after(0, lambda: status_label.config(
                            text="连接成功！API Key 有效", fg='#27ae60'))
                    else:
                        msg = resp.json().get("error", {}).get("message", resp.text[:100])
                        dialog.after(0, lambda m=msg: status_label.config(
                            text=f"连接失败: {m}", fg='#e74c3c'))
                except ImportError:
                    dialog.after(0, lambda: status_label.config(
                        text="缺少 requests 库，请先安装", fg='#e74c3c'))
                except Exception as e:
                    dialog.after(0, lambda m=str(e)[:60]: status_label.config(
                        text=f"连接失败: {m}", fg='#e74c3c'))

            threading.Thread(target=_do_test, daemon=True).start()

        def save_settings():
            api_key = key_var.get().strip()
            self._save_api_key(api_key)
            status_label.config(text="设置已保存！", fg='#27ae60')

        tk.Button(
            btn_frame,
            text="测试连接",
            font=('Microsoft YaHei UI', 9),
            bg='#f5f5f5', fg='#333',
            relief='solid', borderwidth=1,
            cursor='hand2',
            command=test_connection
        ).pack(side=tk.LEFT, padx=(0, 8))

        tk.Button(
            btn_frame,
            text="保存",
            font=('Microsoft YaHei UI', 9),
            bg=self.accent_color, fg='white',
            activebackground='#1a5cc7',
            activeforeground='white',
            relief='flat',
            cursor='hand2',
            width=10,
            command=save_settings
        ).pack(side=tk.RIGHT)

        tk.Button(
            btn_frame,
            text="取消",
            font=('Microsoft YaHei UI', 9),
            bg='#f5f5f5', fg='#333',
            relief='solid', borderwidth=1,
            cursor='hand2',
            command=dialog.destroy
        ).pack(side=tk.RIGHT, padx=(0, 8))

    def _read_api_key(self) -> str:
        """从 .env 文件读取 API Key"""
        env_path = self.base_path / ".env"
        if not env_path.exists():
            return ""
        try:
            with open(env_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("DEEPSEEK_API_KEY=") and not line.startswith("#"):
                        return line.split("=", 1)[1].strip()
        except OSError:
            pass
        return ""

    def _save_api_key(self, api_key: str):
        """将 API Key 保存到 .env 文件"""
        env_path = self.base_path / ".env"
        lines = []

        if env_path.exists():
            try:
                with open(env_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
            except OSError:
                lines = []

        # 查找并替换已有的 DEEPSEEK_API_KEY 行
        found = False
        new_lines = []
        for line in lines:
            if line.strip().startswith("DEEPSEEK_API_KEY=") and not line.strip().startswith("#"):
                new_lines.append(f"DEEPSEEK_API_KEY={api_key}\n")
                found = True
            else:
                new_lines.append(line)

        if not found:
            new_lines.append(f"DEEPSEEK_API_KEY={api_key}\n")

        with open(env_path, 'w', encoding='utf-8') as f:
            f.writelines(new_lines)

    def open_settings(self):
        """配置 DeepSeek 或 OpenAI 兼容接口，并管理可选模型。"""
        if not LLM_AVAILABLE:
            messagebox.showerror("设置不可用", "LLM 配置模块未能加载，请先运行一键安装。")
            return

        settings = load_llm_settings(self.base_path / ".env")
        dialog = tk.Toplevel(self.root)
        dialog.title("LLM 接口与模型设置")
        dialog.geometry("760x650")
        dialog.minsize(700, 600)
        dialog.configure(bg=self.bg_color)
        dialog.transient(self.root)
        dialog.grab_set()

        provider_var = tk.StringVar(value=settings.provider)
        deepseek_key_var = tk.StringVar(value=settings.deepseek_api_key)
        deepseek_model_var = tk.StringVar(value=settings.deepseek_model)
        custom_url_var = tk.StringVar(value=settings.custom_api_base_url)
        custom_key_var = tk.StringVar(value=settings.custom_api_key)
        custom_default_var = tk.StringVar(value=settings.custom_api_model)
        status_var = tk.StringVar(value="设置仅保存在项目根目录的 .env 文件中。")

        outer = ttk.Frame(dialog, padding=16)
        outer.pack(fill=tk.BOTH, expand=True)
        ttk.Label(
            outer,
            text="LLM 接口与模型",
            font=('Microsoft YaHei UI', 13, 'bold'),
        ).pack(anchor=tk.W)
        ttk.Label(
            outer,
            text="选择接口后保存；API Key 不会显示在日志或恢复清单中。",
            style='Muted.TLabel',
        ).pack(anchor=tk.W, pady=(2, 12))

        provider_frame = ttk.Labelframe(
            outer, text="当前接口", padding=10, bootstyle='primary'
        )
        provider_frame.pack(fill=tk.X)
        deepseek_provider_radio = ttk.Radiobutton(
            provider_frame,
            text="DeepSeek",
            value="deepseek",
            variable=provider_var,
            bootstyle='primary',
        )
        deepseek_provider_radio.pack(side=tk.LEFT, padx=(0, 24))
        custom_provider_radio = ttk.Radiobutton(
            provider_frame,
            text="自定义 OpenAI 兼容接口",
            value="custom",
            variable=provider_var,
            bootstyle='primary',
        )
        custom_provider_radio.pack(side=tk.LEFT)

        notebook = ttk.Notebook(outer, bootstyle='primary')

        deepseek_page = ttk.Frame(notebook, padding=14)
        custom_page = ttk.Frame(notebook, padding=14)
        notebook.add(deepseek_page, text="DeepSeek")
        notebook.add(custom_page, text="自定义接口")

        def select_tab_for_provider():
            """单选接口时同步显示对应设置页。"""
            notebook.select(
                custom_page if provider_var.get() == "custom" else deepseek_page
            )

        def select_provider_for_tab(_event=None):
            """切换设置页时同步实际保存的接口类型。"""
            provider_var.set(
                "custom" if notebook.index(notebook.select()) == 1 else "deepseek"
            )

        deepseek_provider_radio.configure(command=select_tab_for_provider)
        custom_provider_radio.configure(command=select_tab_for_provider)
        notebook.bind("<<NotebookTabChanged>>", select_provider_for_tab)

        deepseek_page.columnconfigure(1, weight=1)
        ttk.Label(deepseek_page, text="API Key").grid(
            row=0, column=0, sticky=tk.W, padx=(0, 10), pady=6
        )
        ttk.Entry(
            deepseek_page, textvariable=deepseek_key_var, show="*", bootstyle='primary'
        ).grid(
            row=0, column=1, sticky=(tk.W, tk.E), pady=6
        )
        ttk.Label(deepseek_page, text="模型 ID").grid(
            row=1, column=0, sticky=tk.W, padx=(0, 10), pady=6
        )
        ttk.Entry(
            deepseek_page, textvariable=deepseek_model_var, bootstyle='primary'
        ).grid(
            row=1, column=1, sticky=(tk.W, tk.E), pady=6
        )
        ttk.Label(
            deepseek_page,
            text="默认根地址：https://api.deepseek.com/v1",
            style='Muted.TLabel',
        ).grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(4, 14))

        custom_page.columnconfigure(1, weight=1)
        custom_page.rowconfigure(3, weight=1, minsize=120)
        ttk.Label(custom_page, text="API 根地址").grid(
            row=0, column=0, sticky=tk.W, padx=(0, 10), pady=5
        )
        ttk.Entry(
            custom_page, textvariable=custom_url_var, bootstyle='primary'
        ).grid(
            row=0, column=1, sticky=(tk.W, tk.E), pady=5
        )
        ttk.Label(custom_page, text="API Key（可留空）").grid(
            row=1, column=0, sticky=tk.W, padx=(0, 10), pady=5
        )
        ttk.Entry(
            custom_page, textvariable=custom_key_var, show="*", bootstyle='primary'
        ).grid(
            row=1, column=1, sticky=(tk.W, tk.E), pady=5
        )
        ttk.Label(
            custom_page,
            text="填写到 /v1 这一层；程序会访问 /models 和 /chat/completions。",
            style='Muted.TLabel',
        ).grid(row=2, column=0, columnspan=2, sticky=tk.W, pady=(2, 8))

        model_frame = ttk.Labelframe(
            custom_page,
            text="选择主界面使用的模型（单选，普通单击切换）",
            padding=8,
            bootstyle='info',
        )
        model_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.N, tk.S, tk.W, tk.E))
        model_frame.configure(height=120)
        model_frame.grid_propagate(False)
        model_frame.columnconfigure(0, weight=1)
        model_frame.rowconfigure(0, weight=1)
        model_choices = ttk.ScrolledFrame(
            model_frame,
            height=120,
            auto_hide=False,
            bootstyle='light',
        )
        model_choices.grid(row=0, column=0, sticky=(tk.N, tk.S, tk.W, tk.E))

        available_custom_models = list(dict.fromkeys(settings.custom_api_models))
        if (
            settings.custom_api_model
            and settings.custom_api_model not in available_custom_models
        ):
            available_custom_models.append(settings.custom_api_model)

        current_model_frame = ttk.Frame(custom_page)
        current_model_frame.grid(
            row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(8, 0)
        )
        ttk.Label(current_model_frame, text="当前选择").pack(
            side=tk.LEFT, padx=(0, 10)
        )
        ttk.Label(
            current_model_frame,
            textvariable=custom_default_var,
            style='SectionTitle.TLabel',
        ).pack(side=tk.LEFT, fill=tk.X, expand=True)

        custom_test_status_var = tk.StringVar(
            value="尚未测试。测试会发送一个很小的真实请求，可能产生少量费用。"
        )
        custom_test_status_label = tk.Label(
            custom_page,
            textvariable=custom_test_status_var,
            anchor=tk.NW,
            justify=tk.LEFT,
            height=2,
            wraplength=650,
            fg="#555555",
            bg=self.bg_color,
        )
        custom_test_status_label.grid(
            row=6,
            column=0,
            columnspan=2,
            sticky=(tk.W, tk.E),
            pady=(8, 0),
        )

        def selected_custom_models() -> List[str]:
            selected = custom_default_var.get().strip()
            if selected and selected in available_custom_models:
                return [selected]
            return []

        def on_custom_model_selected():
            provider_var.set("custom")
            selected = custom_default_var.get().strip()
            if selected:
                set_status(f"已选择模型：{selected}；点击保存后应用。", "#16823b")

        def render_model_choices(models: List[str], preferred: str = ""):
            nonlocal available_custom_models
            available_custom_models = list(dict.fromkeys(models))
            for child in model_choices.winfo_children():
                child.destroy()

            if not available_custom_models:
                custom_default_var.set("")
                ttk.Label(
                    model_choices,
                    text="尚未读取到模型。",
                    style='Muted.TLabel',
                ).pack(anchor=tk.W, padx=6, pady=6)
                return

            selected = preferred if preferred in available_custom_models else ""
            if not selected and custom_default_var.get() in available_custom_models:
                selected = custom_default_var.get()
            if not selected:
                selected = available_custom_models[0]
            custom_default_var.set(selected)

            for model_id in available_custom_models:
                ttk.Radiobutton(
                    model_choices,
                    text=model_id,
                    variable=custom_default_var,
                    value=model_id,
                    command=on_custom_model_selected,
                    bootstyle='primary',
                ).pack(fill=tk.X, anchor=tk.W, padx=6, pady=3)

        render_model_choices(available_custom_models, settings.custom_api_model)

        status_label = tk.Label(
            outer,
            textvariable=status_var,
            fg="#555555",
            anchor=tk.NW,
            justify=tk.LEFT,
            height=3,
            wraplength=710,
            bg=self.bg_color,
        )

        def set_status(text: str, color: str = "#555555"):
            if dialog.winfo_exists():
                status_var.set(text)
                status_label.configure(foreground=color)

        def run_in_background(work, success_prefix: str = ""):
            set_status("正在连接，请稍候……", "#b26a00")

            events = queue.Queue()

            def runner():
                try:
                    events.put(("ok", work()))
                except Exception as exc:
                    events.put(("error", str(exc)))

            def poll_result():
                if not dialog.winfo_exists():
                    return
                try:
                    kind, payload = events.get_nowait()
                except queue.Empty:
                    dialog.after(100, poll_result)
                    return
                if kind == "ok":
                    message = (
                        f"{success_prefix}{payload}" if success_prefix else str(payload)
                    )
                    set_status(message, "#16823b")
                else:
                    set_status(f"连接失败：{payload}", "#b42318")

            threading.Thread(target=runner, daemon=True).start()
            dialog.after(100, poll_result)

        def fetch_custom_models():
            provider_var.set("custom")
            try:
                api_base = normalize_api_base_url(custom_url_var.get())
            except ValueError as exc:
                set_status(str(exc), "#b42318")
                return
            api_key = custom_key_var.get().strip()
            set_status("正在读取 /models，请稍候……", "#b26a00")

            def runner():
                try:
                    models = discover_models(api_base, api_key)
                    if not models:
                        raise ValueError("接口没有返回可用模型")

                    def apply_models():
                        previous = custom_default_var.get().strip()
                        render_model_choices(models, previous)
                        set_status(
                            f"已发现 {len(models)} 个模型；请单击选择一个模型。",
                            "#16823b",
                        )

                    self._post_ui(apply_models)
                except Exception as exc:
                    error = str(exc)
                    self._post_ui(
                        set_status,
                        f"读取模型失败：{error}",
                        "#b42318",
                    )

            threading.Thread(target=runner, daemon=True).start()

        def test_deepseek():
            provider_var.set("deepseek")
            model = deepseek_model_var.get().strip() or DEFAULT_DEEPSEEK_MODEL
            api_key = deepseek_key_var.get().strip()
            run_in_background(
                lambda: format_speed_result(test_model_speed(
                    DEEPSEEK_API_BASE,
                    api_key,
                    model,
                )),
                "测试完成；请点击“保存并启用 DeepSeek”。\n",
            )

        def test_custom():
            provider_var.set("custom")
            try:
                api_base = normalize_api_base_url(custom_url_var.get())
            except ValueError as exc:
                set_status(str(exc), "#b42318")
                return
            models = selected_custom_models()
            if not models:
                custom_test_status_var.set("请先读取并选择一个模型。")
                custom_test_status_label.configure(fg="#b42318")
                messagebox.showwarning(
                    "没有可测试模型",
                    "请先点击“读取模型”，再单击选择一个模型。",
                    parent=dialog,
                )
                return

            api_key = custom_key_var.get().strip()
            total = len(models)
            events = queue.Queue()
            custom_test_button.configure(state=tk.DISABLED)
            custom_test_status_label.configure(fg="#b26a00")
            custom_test_status_var.set(f"准备测试 {total} 个模型……")

            def worker():
                results = []
                try:
                    for index, model in enumerate(models, start=1):
                        events.put(("progress", index, total, model))
                        result = test_model_speed(
                            api_base,
                            api_key,
                            model,
                            timeout=30.0,
                        )
                        results.append(result)
                        events.put(("result", index, format_speed_result(result)))
                    events.put(("done", results))
                except Exception as exc:
                    events.put(("fatal", str(exc)))

            completed_lines = []

            def poll_test_events():
                if not dialog.winfo_exists():
                    return
                finished = False
                try:
                    while True:
                        event = events.get_nowait()
                        kind = event[0]
                        if kind == "progress":
                            _, index, count, model = event
                            custom_test_status_var.set(
                                f"正在测试 {index}/{count}：{model}\n"
                                "单个模型最长等待 30 秒，请勿重复点击。"
                            )
                            custom_test_status_label.configure(fg="#b26a00")
                        elif kind == "result":
                            _, _index, line = event
                            completed_lines.append(line)
                            custom_test_status_var.set("\n".join(completed_lines[-3:]))
                        elif kind == "done":
                            results = event[1]
                            success_count = sum(1 for item in results if item.get("ok"))
                            summary = "\n".join(
                                format_speed_result(item) for item in results
                            )
                            custom_test_status_var.set(
                                f"测试完成：{success_count}/{len(results)} 个模型成功。\n"
                                "点击“保存并启用自定义接口”以应用设置。"
                            )
                            custom_test_status_label.configure(
                                fg="#16823b" if success_count else "#b42318"
                            )
                            set_status(
                                "模型测试已完成；测试本身不会自动保存设置。",
                                "#16823b" if success_count else "#b42318",
                            )
                            messagebox.showinfo(
                                "模型测试完成",
                                f"成功：{success_count}/{len(results)}\n\n{summary}",
                                parent=dialog,
                            )
                            finished = True
                        elif kind == "fatal":
                            error = event[1]
                            custom_test_status_var.set(f"测试失败：{error}")
                            custom_test_status_label.configure(fg="#b42318")
                            messagebox.showerror("测试失败", error, parent=dialog)
                            finished = True
                except queue.Empty:
                    pass

                if finished:
                    custom_test_button.configure(state=tk.NORMAL)
                else:
                    dialog.after(100, poll_test_events)

            threading.Thread(target=worker, daemon=True).start()
            dialog.after(100, poll_test_events)

        deepseek_buttons = ttk.Frame(deepseek_page)
        deepseek_buttons.grid(row=3, column=0, columnspan=2, sticky=tk.W)
        ttk.Button(
            deepseek_buttons,
            text="测试当前模型",
            command=test_deepseek,
            bootstyle='secondary outline',
        ).pack(side=tk.LEFT)

        custom_buttons = ttk.Frame(custom_page)
        custom_buttons.grid(row=5, column=0, columnspan=2, sticky=tk.W, pady=(8, 0))
        ttk.Button(
            custom_buttons,
            text="读取模型",
            command=fetch_custom_models,
            bootstyle='secondary outline',
        ).pack(
            side=tk.LEFT, padx=(0, 8)
        )
        custom_test_button = ttk.Button(
            custom_buttons,
            text="测试当前模型",
            command=test_custom,
            bootstyle='info outline',
        )
        custom_test_button.pack(side=tk.LEFT)

        def save_all_settings():
            provider = provider_var.get()
            selected_models = selected_custom_models()
            custom_url = custom_url_var.get().strip()
            if provider == "custom":
                try:
                    custom_url = normalize_api_base_url(custom_url)
                except ValueError as exc:
                    set_status(str(exc), "#b42318")
                    notebook.select(custom_page)
                    return
                if not selected_models:
                    set_status("自定义接口需要选择一个模型。", "#b42318")
                    notebook.select(custom_page)
                    return

            updated = LLMSettings(
                provider=provider,
                deepseek_api_key=deepseek_key_var.get().strip(),
                deepseek_model=deepseek_model_var.get().strip() or DEFAULT_DEEPSEEK_MODEL,
                custom_api_base_url=custom_url,
                custom_api_key=custom_key_var.get().strip(),
                custom_api_models=selected_models,
                custom_api_model=custom_default_var.get().strip(),
            )
            try:
                save_llm_settings(updated, self.base_path / ".env")
                self.llm_settings = updated
                self._refresh_llm_controls()
            except OSError as exc:
                set_status(f"保存失败：{exc}", "#b42318")
                return
            self.log(f"✅ LLM 设置已保存：{self._llm_provider_display_name()} / {updated.active_model}")
            messagebox.showinfo(
                "保存成功",
                f"已启用：{self._llm_provider_display_name()}\n默认模型：{updated.active_model}",
                parent=dialog,
            )
            dialog.destroy()

        def save_custom_settings():
            provider_var.set("custom")
            notebook.select(custom_page)
            save_all_settings()

        def save_deepseek_settings():
            provider_var.set("deepseek")
            notebook.select(deepseek_page)
            save_all_settings()

        ttk.Button(
            deepseek_buttons,
            text="保存并启用 DeepSeek",
            command=save_deepseek_settings,
            bootstyle='success',
        ).pack(side=tk.LEFT, padx=(16, 0))

        ttk.Button(
            custom_buttons,
            text="保存并启用自定义接口",
            command=save_custom_settings,
            bootstyle='success',
        ).pack(side=tk.LEFT, padx=(16, 0))

        action_frame = ttk.Frame(outer)
        ttk.Button(
            action_frame,
            text="取消",
            command=dialog.destroy,
            bootstyle='secondary outline',
        ).pack(side=tk.RIGHT)
        ttk.Button(
            action_frame,
            text="应用并保存",
            command=save_all_settings,
            bootstyle='primary',
        ).pack(
            side=tk.RIGHT, padx=(0, 8)
        )

        # 先从底部为操作栏和状态栏保留空间，再让 Notebook 填满剩余区域。
        # 这样在 760x650 的默认尺寸下，保存/取消控件不会被内容页挤出窗口。
        action_frame.pack(side=tk.BOTTOM, fill=tk.X)
        status_label.pack(side=tk.BOTTOM, fill=tk.X, pady=(8, 8))
        notebook.pack(fill=tk.BOTH, expand=True, pady=(12, 8))

        if settings.provider == "custom":
            notebook.select(custom_page)


def main():
    """主函数"""
    if ttk is None:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            'PaperMiner 缺少界面依赖',
            '未检测到 ttkbootstrap。\n\n'
            '请先运行 Setup.exe 或“清理重装”，安装 PaperMiner v1.4.5 依赖。\n\n'
            f'详细信息：{_TTKBOOTSTRAP_IMPORT_ERROR}',
            parent=root,
        )
        root.destroy()
        return

    root = ttk.App(
        title=f'{__app_name__} v{__version__}',
        theme='bootstrap-light',
        high_dpi=True,
    )
    BatchPDFProcessorGUI(root)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        print("\n检测到 Ctrl+C，程序已退出。")


if __name__ == "__main__":
    main()
