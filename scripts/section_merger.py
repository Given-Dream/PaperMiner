"""Merge article titles, sections, chart summaries, availability, and references."""
import json
import os
from pathlib import Path
import re


def _article_title(article_dir: Path) -> str:
    """Restore the full source title when path protection shortened a folder."""
    manifest = article_dir / ".paperminer-source.json"
    try:
        payload = json.loads(manifest.read_text(encoding="utf-8-sig"))
    except (OSError, ValueError, TypeError, json.JSONDecodeError):
        return article_dir.name
    source_stem = payload.get("source_stem") if isinstance(payload, dict) else None
    if isinstance(source_stem, str) and source_stem.strip():
        return source_stem.strip()
    return article_dir.name


def find_sections(extract_root: Path):
    """Return {section_name: [(article_title, markdown_path), ...]} for extracted papers."""
    result = {}
    # Some runs produce only the full article Markdown (without Sections/).
    # Materialize standard sections from that source so it can still be merged.
    for article_dir in extract_root.iterdir() if extract_root.exists() else []:
        if not article_dir.is_dir() or article_dir.name == "MergedSections":
            continue
        sections_dir = article_dir / "Sections"
        source_files = [p for p in article_dir.glob("*.md") if "_合并" not in p.stem]
        if source_files:
            # Fill only missing standard sections; preserve explicitly extracted files.
            _materialize_standard_sections(source_files[0], sections_dir, only_missing=True)

    for sections_dir in extract_root.glob("*/Sections"):
        if not sections_dir.is_dir():
            continue
        title = _article_title(sections_dir.parent)
        for path in sorted(sections_dir.glob("*.md")):
            result.setdefault(path.stem, []).append((title, path))
    return result


def find_chart_markdowns(extract_root: Path):
    """Return ``[(article_title, markdown_path), ...]`` from each Word folder."""
    result = []
    article_dirs = sorted(
        (path for path in extract_root.iterdir() if path.is_dir()),
        key=lambda path: path.name.casefold(),
    ) if extract_root.exists() else []
    for article_dir in article_dirs:
        if article_dir.name == "MergedSections":
            continue
        word_dir = article_dir / "Word"
        if not word_dir.is_dir():
            continue
        markdown_files = sorted(
            (
                path for path in word_dir.glob("*.md")
                if "_合并" not in path.stem and path.stat().st_size > 0
            ),
            key=lambda path: path.name.casefold(),
        )
        for path in markdown_files:
            result.append((_article_title(article_dir), path))
    return result


def find_open_source_markdowns(extract_root: Path):
    """Return only reports containing trusted code/data addresses.

    A v2 scan marker with zero results suppresses any stale v1.4.6 audit report
    without deleting user files. Legacy reports without a marker are accepted
    only when they contain at least one address.
    """
    result = []
    article_dirs = sorted(
        (path for path in extract_root.iterdir() if path.is_dir()),
        key=lambda path: path.name.casefold(),
    ) if extract_root.exists() else []
    for article_dir in article_dirs:
        if article_dir.name == "MergedSections":
            continue
        source_dir = article_dir / "OpenSource"
        marker = source_dir / "availability_scan.json"
        source = source_dir / "代码与数据可用性.md"
        if marker.is_file():
            try:
                scan = json.loads(marker.read_text(encoding="utf-8", errors="replace"))
                if int(scan.get("trusted_link_count", 0)) <= 0:
                    continue
            except (OSError, ValueError, TypeError, AttributeError):
                continue
            if source.is_file() and source.stat().st_size > 0:
                content = source.read_text(encoding="utf-8", errors="replace")
                if re.search(r"(?m)^- 地址：", content):
                    result.append((_article_title(article_dir), source))
            continue

        legacy = source_dir / "开源代码地址.md"
        if legacy.is_file() and legacy.stat().st_size > 0:
            content = legacy.read_text(encoding="utf-8", errors="replace")
            if re.search(r"(?m)^- 地址：", content):
                result.append((_article_title(article_dir), legacy))
    return result


def _count_reference_entries(content: str) -> int:
    """Conservatively count ordered bibliography lines in a saved report."""
    return len(
        re.findall(
            r"(?m)^\s*(?:[-*+]\s+)?(?:"
            r"\[\s*[A-Za-z]?\s*\d+\s*\]|"
            r"\(\s*[A-Za-z]?\s*\d+\s*\)|"
            r"[A-Za-z]?\s*\d+\s*[.)．、]"
            r")\s+\S",
            content,
        )
    )


def find_reference_markdowns(extract_root: Path):
    """Return ``(article_title, report_path, entry_count)`` for references.

    A completed zero-result marker suppresses a stale report.  Reports from
    older/custom runs without a marker are accepted only when numbered
    bibliography entries can actually be counted.
    """
    result = []
    article_dirs = sorted(
        (path for path in extract_root.iterdir() if path.is_dir()),
        key=lambda path: path.name.casefold(),
    ) if extract_root.exists() else []
    for article_dir in article_dirs:
        if article_dir.name == "MergedSections":
            continue
        source_dir = article_dir / "References"
        source = source_dir / "参考文献.md"
        marker = source_dir / "references_scan.json"
        if marker.is_file():
            try:
                scan = json.loads(marker.read_text(encoding="utf-8", errors="replace"))
                if not isinstance(scan, dict) or scan.get("scan_completed") is not True:
                    continue
                entry_count = int(scan.get("entry_count", 0))
                report_written = scan.get("report_written") is True
            except (OSError, ValueError, TypeError, AttributeError):
                continue
            if entry_count <= 0 or not report_written:
                continue
            if source.is_file() and source.stat().st_size > 0:
                content = source.read_text(encoding="utf-8", errors="replace")
                actual_entry_count = _count_reference_entries(content)
                if actual_entry_count > 0:
                    result.append(
                        (_article_title(article_dir), source, actual_entry_count)
                    )
            continue

        if source.is_file() and source.stat().st_size > 0:
            content = source.read_text(encoding="utf-8", errors="replace")
            entry_count = _count_reference_entries(content)
            if entry_count > 0:
                result.append((_article_title(article_dir), source, entry_count))
    return result


def find_title_markdowns(extract_root: Path):
    """Return completed per-paper title reports and their audited result.

    Title reports are accepted only when the matching scan marker says the scan
    completed and the report was written.  This avoids merging a stale report
    left by a previous run with different extraction options.
    """
    result = []
    article_dirs = sorted(
        (path for path in extract_root.iterdir() if path.is_dir()),
        key=lambda path: path.name.casefold(),
    ) if extract_root.exists() else []
    for article_dir in article_dirs:
        if article_dir.name == "MergedSections":
            continue
        title_dir = article_dir / "Title"
        marker = title_dir / "title_scan.json"
        report = title_dir / "文章标题.md"
        if not marker.is_file():
            continue
        try:
            scan = json.loads(marker.read_text(encoding="utf-8", errors="replace"))
            if (
                not isinstance(scan, dict)
                or scan.get("schema_version") != 1
                or scan.get("scan_completed") is not True
                or scan.get("report_written") is not True
            ):
                continue
            title = str(scan.get("title") or "").strip()
            confidence = str(scan.get("confidence") or "未标注").strip()
        except (OSError, ValueError, TypeError, AttributeError):
            continue
        if not title or not report.is_file() or report.stat().st_size <= 0:
            continue
        report_text = report.read_text(encoding="utf-8", errors="replace")
        if title not in report_text:
            continue
        result.append((_article_title(article_dir), title, confidence, report))
    return result


_MARKDOWN_LINK_RE = re.compile(
    r"(!?\[(?:[^\[\]]|\[[^\]]*\])*\]\()([^)]+)(\))"
)


def _rewrite_chart_links(
    content: str,
    source: Path,
    target: Path,
    extract_root: Path,
):
    """Keep local images valid after moving Word Markdown into MergedSections."""
    extract_root = extract_root.resolve()
    target_parent = target.parent.resolve()

    def replace(match):
        raw_destination = match.group(2).strip()
        destination = raw_destination
        if destination.startswith("<") and destination.endswith(">"):
            destination = destination[1:-1].strip()
        if (
            not destination
            or destination.startswith(("#", "/", "\\"))
            or re.match(r"^[a-z][a-z0-9+.-]*:", destination, re.IGNORECASE)
            or Path(destination).is_absolute()
        ):
            return match.group(0)

        resolved = (source.parent / destination).resolve()
        try:
            resolved.relative_to(extract_root)
        except ValueError:
            return match.group(0)

        relative = os.path.relpath(resolved, target_parent).replace("\\", "/")
        if any(character.isspace() for character in relative) or any(
            character in relative for character in "()"
        ):
            relative = f"<{relative}>"
        return f"{match.group(1)}{relative}{match.group(3)}"

    return _MARKDOWN_LINK_RE.sub(replace, content)


def _prepare_chart_markdown(
    content: str,
    source: Path,
    target: Path,
    extract_root: Path,
):
    """Remove the source title, nest its headings, and rebase local links."""
    content = content.lstrip("\ufeff")
    lines = content.splitlines()
    for index, line in enumerate(lines):
        if not line.strip():
            continue
        if re.match(r"^#\s+", line):
            del lines[index]
        break

    nested = []
    for line in lines:
        heading = re.match(r"^(#{1,5})(\s+.*)$", line)
        if heading:
            line = f"#{heading.group(1)}{heading.group(2)}"
        nested.append(line)
    prepared = "\n".join(nested).strip()
    return _rewrite_chart_links(prepared, source, target, extract_root)


def _materialize_standard_sections(source: Path, sections_dir: Path, only_missing: bool = False):
    """Extract common paper sections from a full Markdown source file."""
    text = source.read_text(encoding="utf-8", errors="replace")
    headings = list(re.finditer(r"(?m)^(#{1,6})\s+(.+?)\s*$", text))
    if not headings:
        return
    aliases = {
        "abstract": "Abstract",
        "introduction": "Introduction",
        "method": "Methods",
        "methods": "Methods",
        "materials and methods": "Methods",
        "methodology": "Methods",
        "numerical model": "Methods",
        "mathematical model": "Methods",
        "results": "Results & Discussion",
        "results and discussion": "Results & Discussion",
        "results & discussion": "Results & Discussion",
        "discussion": "Results & Discussion",
        "model validation": "Results & Discussion",
        "conclusion": "Conclusion",
        "conclusions": "Conclusion",
    }
    def classify(title):
        normalized_title = re.sub(r"[^a-z]+", " ", title.lower()).strip()
        canonical_title = aliases.get(normalized_title)
        if canonical_title is None and normalized_title.replace(" ", "") == "abstract":
            canonical_title = "Abstract"
        return canonical_title, normalized_title

    extracted = {}
    for index, match in enumerate(headings):
        raw_title = match.group(2).strip().strip("#").strip()
        canonical, normalized = classify(raw_title)
        if not canonical:
            continue
        end = len(text)
        for next_match in headings[index + 1:]:
            next_title = next_match.group(2).strip().strip("#").strip()
            next_canonical, next_normalized = classify(next_title)
            if next_canonical and next_canonical != canonical:
                end = next_match.start()
                break
            if next_normalized in {"acknowledgment", "acknowledgments", "acknowledgement", "acknowledgements", "references", "bibliography"}:
                end = next_match.start()
                break
        content = text[match.end():end].strip()
        if content and canonical not in extracted:
            extracted[canonical] = content
    if extracted:
        sections_dir.mkdir(parents=True, exist_ok=True)
        for name, content in extracted.items():
            target = sections_dir / f"{name}.md"
            if not only_missing or not target.exists():
                target.write_text(content + "\n", encoding="utf-8")


def merge_section_to_docx(extract_root: Path, section_name: str, output_path: Path):
    """Merge one same-named section from all papers into a DOCX."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    matches = find_sections(extract_root).get(section_name, [])
    if not matches:
        raise ValueError(f"未找到章节: {section_name}")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    title = doc.add_paragraph(section_name, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, (article_title, path) in enumerate(matches):
        if index:
            doc.add_page_break()
        doc.add_heading(article_title, level=1)
        text = path.read_text(encoding="utf-8", errors="replace")
        for block in re.split(r"\n\s*\n", text):
            block = block.strip()
            if not block:
                continue
            lines = block.splitlines()
            if len(lines) == 1 and re.match(r"^#{1,6}\s+", lines[0]):
                level = min(len(lines[0]) - len(lines[0].lstrip("#")), 3)
                doc.add_heading(re.sub(r"^#{1,6}\s+", "", lines[0]).strip(), level=level)
            else:
                p = doc.add_paragraph()
                p.add_run("\n".join(lines))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path, len(matches)


def merge_all_sections_to_docx(extract_root: Path, output_path: Path):
    """Merge every same-named section into a separate DOCX file."""
    grouped = find_sections(extract_root)
    if not grouped:
        raise ValueError("未找到可合并的章节")
    output_path.mkdir(parents=True, exist_ok=True)
    outputs = []
    total = 0
    for section_name in sorted(grouped):
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", section_name).strip() or "未命名章节"
        output, count = merge_section_to_docx(extract_root, section_name, output_path / f"{safe_name}_合并.docx")
        outputs.append(output)
        total += count
    return outputs, total, len(grouped)


def _write_section_markdowns(grouped, output_path: Path):
    """Write pre-grouped sections and return outputs plus source count."""
    output_path.mkdir(parents=True, exist_ok=True)
    outputs = []
    total = 0
    for section_name in sorted(grouped):
        safe_name = re.sub(r'[\\/:*?"<>|]', "_", section_name).strip() or "未命名章节"
        target = output_path / f"{safe_name}_合并.md"
        chunks = [f"# {section_name}\n"]
        for article_title, source in grouped[section_name]:
            content = source.read_text(encoding="utf-8", errors="replace").strip()
            chunks.append(f"\n## 【{article_title}】\n\n{content}\n")
            total += 1
        target.write_text("\n".join(chunks), encoding="utf-8")
        outputs.append(target)
    return outputs, total


def merge_all_sections_to_markdown(extract_root: Path, output_path: Path):
    """Write one Markdown file per section, preserving source LaTeX verbatim."""
    grouped = find_sections(extract_root)
    if not grouped:
        raise ValueError("未找到可合并的章节")
    outputs, total = _write_section_markdowns(grouped, output_path)
    return outputs, total, len(grouped)


def merge_chart_markdowns(extract_root: Path, target: Path, sources=None):
    """Merge Word-folder chart Markdown and keep all relative media links valid."""
    sources = find_chart_markdowns(extract_root) if sources is None else sources
    if not sources:
        raise ValueError("未找到 Word 文件夹中的图表 Markdown")

    target.parent.mkdir(parents=True, exist_ok=True)
    chunks = [
        "# 图表汇总\n",
        "> 来源：各论文 Word 文件夹中的 Markdown 图表汇总。\n",
    ]
    for article_title, source in sources:
        content = source.read_text(encoding="utf-8", errors="replace")
        content = _prepare_chart_markdown(content, source, target, extract_root)
        chunks.append(f"\n## 【{article_title}】\n")
        if content:
            chunks.append(f"\n{content}\n")
        else:
            chunks.append("\n> 此图表 Markdown 为空。\n")

    target.write_text("\n".join(chunks).rstrip() + "\n", encoding="utf-8")
    return target, len(sources)


def merge_open_source_markdowns(extract_root: Path, target: Path, sources=None):
    """Merge per-paper code/data availability reports."""
    sources = find_open_source_markdowns(extract_root) if sources is None else sources
    if not sources:
        raise ValueError("未找到含可信地址的代码/数据可用性 Markdown")

    target.parent.mkdir(parents=True, exist_ok=True)
    chunks = [
        "# 代码与数据可用性汇总\n",
        "> 来源：各论文 OpenSource/代码与数据可用性.md；DOI 不收录，结果需结合原文人工核验。\n",
    ]
    total_links = 0
    for article_title, source in sources:
        content = source.read_text(encoding="utf-8", errors="replace")
        total_links += len(re.findall(r"(?m)^- 地址：", content))
        content = _prepare_chart_markdown(content, source, target, extract_root)
        chunks.append(f"\n## 【{article_title}】\n")
        if content:
            chunks.append(f"\n{content}\n")
        else:
            chunks.append("\n> 此代码/数据可用性报告为空。\n")

    chunks.insert(
        2,
        f"> 共汇总 {len(sources)} 篇论文，识别到 {total_links} 个候选地址。\n",
    )
    target.write_text("\n".join(chunks).rstrip() + "\n", encoding="utf-8")
    return target, len(sources), total_links


def merge_reference_markdowns(extract_root: Path, target: Path, sources=None):
    """Merge per-paper reference reports without deduplication or rewriting."""
    sources = find_reference_markdowns(extract_root) if sources is None else sources
    if not sources:
        raise ValueError("未找到含参考文献条目的 Markdown")

    target.parent.mkdir(parents=True, exist_ok=True)
    total_entries = sum(max(0, int(entry_count)) for _, _, entry_count in sources)
    chunks = [
        "# 参考文献汇总\n",
        "> 来源：各论文 References/参考文献.md。按论文分组保留原顺序和原编号，不跨论文去重或改写。\n",
        f"> 共汇总 {len(sources)} 篇论文，累计 {total_entries} 条参考文献；正式引用前请与原 PDF 核对。\n",
    ]
    for article_title, source, _entry_count in sources:
        content = source.read_text(encoding="utf-8", errors="replace")
        content = _prepare_chart_markdown(content, source, target, extract_root)
        chunks.append(f"\n## 【{article_title}】\n")
        if content:
            chunks.append(f"\n{content}\n")
        else:
            chunks.append("\n> 此参考文献报告为空。\n")

    target.write_text("\n".join(chunks).rstrip() + "\n", encoding="utf-8")
    return target, len(sources), total_entries


def merge_title_markdowns(extract_root: Path, target: Path, sources=None):
    """Merge audited article titles without silently deduplicating papers."""
    sources = find_title_markdowns(extract_root) if sources is None else sources
    if not sources:
        raise ValueError("未找到已完成的文章标题 Markdown")

    target.parent.mkdir(parents=True, exist_ok=True)
    review_count = sum(confidence in {"需核查", "未识别"} for _, _, confidence, _ in sources)
    chunks = [
        "# 文章标题汇总\n",
        "> 来源：各论文 Title/文章标题.md；按输入论文逐篇列出，不跨论文去重。\n",
        f"> 共汇总 {len(sources)} 篇论文；其中 {review_count} 篇需要人工核查原 PDF 首页。\n",
    ]
    for index, (source_stem, title, confidence, _report) in enumerate(sources, start=1):
        chunks.extend(
            (
                f"\n## {index}. {title}\n",
                f"\n- 来源文件：{source_stem}.pdf\n",
                f"- 置信度：{confidence}\n",
            )
        )
    target.write_text("\n".join(chunks).rstrip() + "\n", encoding="utf-8")
    return target, len(sources), review_count


def merge_all_sections_and_charts_to_markdown(
    extract_root: Path,
    output_path: Path,
):
    """Merge same-named sections plus Word-folder chart Markdown in one action."""
    grouped = find_sections(extract_root)
    chart_sources = find_chart_markdowns(extract_root)
    if not grouped and not chart_sources:
        raise ValueError("未找到可合并的章节或 Word 文件夹图表 Markdown")

    output_path.mkdir(parents=True, exist_ok=True)
    outputs = []
    total_section_articles = 0
    if grouped:
        section_outputs, total_section_articles = _write_section_markdowns(
            grouped,
            output_path,
        )
        outputs.extend(section_outputs)

    chart_count = 0
    if chart_sources:
        chart_output, chart_count = merge_chart_markdowns(
            extract_root,
            output_path / "图表汇总_合并.md",
            chart_sources,
        )
        outputs.append(chart_output)

    return outputs, total_section_articles, len(grouped), chart_count


def merge_all_sections_charts_code_and_references_to_markdown(
    extract_root: Path,
    output_path: Path,
):
    """Merge titles, sections, charts, code/data reports, and references."""
    title_sources = find_title_markdowns(extract_root)
    grouped = find_sections(extract_root)
    chart_sources = find_chart_markdowns(extract_root)
    code_sources = find_open_source_markdowns(extract_root)
    reference_sources = find_reference_markdowns(extract_root)
    if (
        not title_sources
        and not grouped
        and not chart_sources
        and not code_sources
        and not reference_sources
    ):
        raise ValueError(
            "未找到可合并的文章标题、章节、图表、代码/数据可用性或参考文献 Markdown"
        )

    output_path.mkdir(parents=True, exist_ok=True)
    outputs = []
    title_source_count = 0
    title_review_count = 0
    if title_sources:
        title_output, title_source_count, title_review_count = merge_title_markdowns(
            extract_root,
            output_path / "文章标题_合并.md",
            title_sources,
        )
        outputs.append(title_output)

    total_section_articles = 0
    if grouped:
        section_outputs, total_section_articles = _write_section_markdowns(
            grouped,
            output_path,
        )
        outputs.extend(section_outputs)

    chart_count = 0
    if chart_sources:
        chart_output, chart_count = merge_chart_markdowns(
            extract_root,
            output_path / "图表汇总_合并.md",
            chart_sources,
        )
        outputs.append(chart_output)

    code_source_count = 0
    code_link_count = 0
    if code_sources:
        code_output, code_source_count, code_link_count = merge_open_source_markdowns(
            extract_root,
            output_path / "代码与数据可用性_合并.md",
            code_sources,
        )
        outputs.append(code_output)

    reference_source_count = 0
    reference_entry_count = 0
    if reference_sources:
        (
            reference_output,
            reference_source_count,
            reference_entry_count,
        ) = merge_reference_markdowns(
            extract_root,
            output_path / "参考文献_合并.md",
            reference_sources,
        )
        outputs.append(reference_output)

    return (
        outputs,
        total_section_articles,
        len(grouped),
        chart_count,
        code_source_count,
        code_link_count,
        reference_source_count,
        reference_entry_count,
        title_source_count,
        title_review_count,
    )


def merge_all_sections_charts_and_code_to_markdown(
    extract_root: Path,
    output_path: Path,
):
    """Backward-compatible six-field wrapper for pre-v1.4.21 callers.

    The reference Markdown is still generated as part of the operation, while
    the historical return shape remains unchanged for external integrations.
    """
    result = merge_all_sections_charts_code_and_references_to_markdown(
        extract_root,
        output_path,
    )
    return result[:6]


def _legacy_merge_all_sections_to_docx(extract_root: Path, output_path: Path):
    """Legacy single-document implementation retained for compatibility."""
    from docx import Document
    from docx.shared import Cm, Pt
    grouped = find_sections(extract_root)
    if not grouped:
        raise ValueError("未找到可合并的章节")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(2.54)
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_paragraph("合并章节", style="Title")
    total = 0
    for section_name in sorted(grouped):
        doc.add_heading(section_name, level=1)
        for article_title, path in grouped[section_name]:
            doc.add_heading(article_title, level=2)
            for block in re.split(r"\n\s*\n", path.read_text(encoding="utf-8", errors="replace")):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)
            total += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path, total, len(grouped)
